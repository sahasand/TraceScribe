"""ICF (Informed Consent Form) generator."""

import json
import logging
from typing import Optional
from dataclasses import dataclass, field
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .base import BaseDocumentGenerator
from app.modules.ai.prompts.icf_generation import ICF_GENERATION_PROMPT, ICF_POLISH_PROMPT

logger = logging.getLogger(__name__)


@dataclass
class ICFContent:
    """ICF content structure."""
    study_title: str = ""
    protocol_number: str = ""
    sponsor: str = ""
    phase: str = ""
    study_purpose: str = ""
    procedures_section: str = ""
    time_commitment: str = ""
    risks_section: str = ""
    benefits_section: str = ""
    alternatives_section: str = ""
    confidentiality_section: str = ""
    compensation_section: str = ""
    voluntary_section: str = ""
    visits: list = field(default_factory=list)
    adverse_events: list = field(default_factory=list)


class ICFGenerator(BaseDocumentGenerator[ICFContent]):
    """Generator for Informed Consent Forms."""

    template_name = "standard_icf.docx"
    document_type = "icf"
    requires_polish = True

    async def extract_for_document(self, protocol_data: dict) -> ICFContent:
        """Extract ICF-relevant data from protocol."""
        metadata = protocol_data.get("metadata", {})
        design = protocol_data.get("design", {})
        visits = protocol_data.get("visits", [])
        adverse_events = protocol_data.get("adverse_events", [])
        procedures = protocol_data.get("procedures", [])
        eligibility = protocol_data.get("eligibility", {})
        ip = protocol_data.get("investigational_product", {}) or {}

        # Generate ICF content using Gemini
        if self.gemini:
            generated = await self._generate_icf_content(protocol_data)
        else:
            generated = self._generate_fallback_content(protocol_data)

        return ICFContent(
            study_title=metadata.get("title", "Clinical Study"),
            protocol_number=metadata.get("protocol_number", ""),
            sponsor=metadata.get("sponsor", ""),
            phase=metadata.get("phase", ""),
            study_purpose=generated.get("study_purpose", ""),
            procedures_section=generated.get("procedures_section", ""),
            time_commitment=generated.get("time_commitment", ""),
            risks_section=generated.get("risks_section", ""),
            benefits_section=generated.get("benefits_section", ""),
            alternatives_section=generated.get("alternatives_section", ""),
            confidentiality_section=generated.get("confidentiality_section", ""),
            compensation_section=generated.get("compensation_section", ""),
            voluntary_section=generated.get("voluntary_section", ""),
            visits=visits,
            adverse_events=adverse_events,
        )

    async def _generate_icf_content(self, protocol_data: dict) -> dict:
        """Generate ICF content using Gemini."""
        prompt = ICF_GENERATION_PROMPT.format(
            protocol_data=json.dumps(protocol_data, indent=2)
        )

        try:
            response = await self.gemini.generate(prompt, temperature=0.3)

            # Parse JSON response
            json_str = response.strip()
            if json_str.startswith("```"):
                json_str = json_str.split("```")[1]
                if json_str.startswith("json"):
                    json_str = json_str[4:]
            json_str = json_str.strip()

            return json.loads(json_str)
        except Exception as e:
            logger.error(f"Failed to generate ICF content: {e}")
            return self._generate_fallback_content(protocol_data)

    def _generate_fallback_content(self, protocol_data: dict) -> dict:
        """Generate fallback ICF content without AI."""
        metadata = protocol_data.get("metadata", {})
        design = protocol_data.get("design", {})
        ip = protocol_data.get("investigational_product", {}) or {}

        return {
            "study_purpose": f"This research study is being done to learn about {ip.get('name', 'an investigational treatment')} for {metadata.get('indication', 'your condition')}. The study will help researchers understand if the treatment is safe and effective.",
            "procedures_section": "During this study, you will have regular visits with the study team. At these visits, you may have physical exams, blood tests, and other procedures to check your health and how the treatment is working.",
            "time_commitment": f"This study will last approximately {design.get('study_duration_weeks', 'several')} weeks. You will need to attend study visits as scheduled by your study team.",
            "risks_section": "Like all treatments, this study treatment may cause side effects. The most common side effects and their frequency are listed in the following sections.",
            "benefits_section": "You may or may not benefit from being in this study. The information learned from this study may help other people in the future.",
            "alternatives_section": "Instead of being in this study, you could choose to receive standard treatment or no treatment. Your study doctor can discuss these options with you.",
            "confidentiality_section": "Your medical records will be kept confidential. Your name will not appear in any published results. Only authorized personnel will have access to your information.",
            "compensation_section": "If you are injured as a result of taking part in this study, medical treatment is available. Please contact the study team immediately if you experience any problems.",
            "voluntary_section": "Taking part in this study is your choice. You may decide not to take part or may leave the study at any time. Your decision will not affect your regular medical care.",
        }

    async def polish_content(self, extracted: ICFContent) -> ICFContent:
        """Polish ICF content using Claude."""
        if not self.claude:
            return extracted

        content_dict = {
            "study_purpose": extracted.study_purpose,
            "procedures_section": extracted.procedures_section,
            "time_commitment": extracted.time_commitment,
            "risks_section": extracted.risks_section,
            "benefits_section": extracted.benefits_section,
            "alternatives_section": extracted.alternatives_section,
            "confidentiality_section": extracted.confidentiality_section,
            "compensation_section": extracted.compensation_section,
            "voluntary_section": extracted.voluntary_section,
        }

        try:
            polished = await self.claude.polish_regulatory_text(
                content=json.dumps(content_dict),
                document_type="ICF",
                guidelines="6-8th grade reading level, plain language, no jargon"
            )

            polished_dict = json.loads(polished)

            extracted.study_purpose = polished_dict.get("study_purpose", extracted.study_purpose)
            extracted.procedures_section = polished_dict.get("procedures_section", extracted.procedures_section)
            extracted.time_commitment = polished_dict.get("time_commitment", extracted.time_commitment)
            extracted.risks_section = polished_dict.get("risks_section", extracted.risks_section)
            extracted.benefits_section = polished_dict.get("benefits_section", extracted.benefits_section)
            extracted.alternatives_section = polished_dict.get("alternatives_section", extracted.alternatives_section)
            extracted.confidentiality_section = polished_dict.get("confidentiality_section", extracted.confidentiality_section)
            extracted.compensation_section = polished_dict.get("compensation_section", extracted.compensation_section)
            extracted.voluntary_section = polished_dict.get("voluntary_section", extracted.voluntary_section)

        except Exception as e:
            logger.warning(f"Failed to polish ICF content: {e}")

        return extracted

    def build_template_context(self, extracted: ICFContent) -> dict:
        """Build context for ICF template."""
        # Organize adverse events by frequency
        ae_by_frequency = {
            "very_common": [],
            "common": [],
            "uncommon": [],
            "rare": [],
            "unknown": [],
        }

        for ae in extracted.adverse_events:
            freq = ae.get("frequency", "Unknown").lower()
            plain_lang = ae.get("plain_language") or ae.get("term", "")

            if "very common" in freq or ">10%" in freq:
                ae_by_frequency["very_common"].append(plain_lang)
            elif "common" in freq or "1-10%" in freq:
                ae_by_frequency["common"].append(plain_lang)
            elif "uncommon" in freq or "<1%" in freq:
                ae_by_frequency["uncommon"].append(plain_lang)
            elif "rare" in freq:
                ae_by_frequency["rare"].append(plain_lang)
            else:
                ae_by_frequency["unknown"].append(plain_lang)

        return {
            "title": f"INFORMED CONSENT FORM",
            "study_title": extracted.study_title,
            "protocol_number": extracted.protocol_number,
            "sponsor": extracted.sponsor,
            "phase": extracted.phase,
            "purpose_section": extracted.study_purpose,
            "procedures_section": extracted.procedures_section,
            "time_commitment": extracted.time_commitment,
            "risks_section": extracted.risks_section,
            "benefits_section": extracted.benefits_section,
            "alternatives_section": extracted.alternatives_section,
            "confidentiality_section": extracted.confidentiality_section,
            "compensation_section": extracted.compensation_section,
            "voluntary_section": extracted.voluntary_section,
            "visits": extracted.visits,
            "adverse_events": extracted.adverse_events,
            "ae_very_common": ae_by_frequency["very_common"],
            "ae_common": ae_by_frequency["common"],
            "ae_uncommon": ae_by_frequency["uncommon"],
            "ae_rare": ae_by_frequency["rare"],
            "ae_unknown": ae_by_frequency["unknown"],
            # Placeholders for site-specific info
            "site_name": "{{site_name}}",
            "pi_name": "{{pi_name}}",
            "pi_phone": "{{pi_phone}}",
            "irb_name": "{{irb_name}}",
            "irb_phone": "{{irb_phone}}",
            "emergency_contact": "{{emergency_contact}}",
        }

    async def _add_content_to_document(self, doc: Document, context: dict) -> None:
        """Add ICF content to document."""
        # Title
        title = doc.add_heading("INFORMED CONSENT FORM", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Study info
        doc.add_paragraph(f"Study Title: {context.get('study_title', '')}")
        doc.add_paragraph(f"Protocol Number: {context.get('protocol_number', '')}")
        doc.add_paragraph(f"Sponsor: {context.get('sponsor', '')}")
        doc.add_paragraph("")

        # Sections
        sections = [
            ("1. PURPOSE OF THE STUDY", context.get("purpose_section", "")),
            ("2. PROCEDURES", context.get("procedures_section", "")),
            ("3. TIME COMMITMENT", context.get("time_commitment", "")),
            ("4. RISKS AND SIDE EFFECTS", context.get("risks_section", "")),
            ("5. BENEFITS", context.get("benefits_section", "")),
            ("6. ALTERNATIVES", context.get("alternatives_section", "")),
            ("7. CONFIDENTIALITY", context.get("confidentiality_section", "")),
            ("8. COMPENSATION FOR INJURY", context.get("compensation_section", "")),
            ("9. VOLUNTARY PARTICIPATION", context.get("voluntary_section", "")),
        ]

        for heading, content in sections:
            doc.add_heading(heading, level=1)
            doc.add_paragraph(content)
            doc.add_paragraph("")

        # Add risks by frequency if available
        if any([
            context.get("ae_very_common"),
            context.get("ae_common"),
            context.get("ae_uncommon"),
            context.get("ae_rare"),
        ]):
            doc.add_heading("DETAILED SIDE EFFECTS BY FREQUENCY", level=2)

            if context.get("ae_very_common"):
                doc.add_heading("Very Common (more than 1 in 10 people):", level=3)
                for ae in context["ae_very_common"]:
                    doc.add_paragraph(f"• {ae}", style="List Bullet")

            if context.get("ae_common"):
                doc.add_heading("Common (1 to 10 in 100 people):", level=3)
                for ae in context["ae_common"]:
                    doc.add_paragraph(f"• {ae}", style="List Bullet")

            if context.get("ae_uncommon"):
                doc.add_heading("Uncommon (less than 1 in 100 people):", level=3)
                for ae in context["ae_uncommon"]:
                    doc.add_paragraph(f"• {ae}", style="List Bullet")

            if context.get("ae_rare"):
                doc.add_heading("Rare:", level=3)
                for ae in context["ae_rare"]:
                    doc.add_paragraph(f"• {ae}", style="List Bullet")

        # Contact section
        doc.add_heading("10. CONTACTS", level=1)
        doc.add_paragraph("For questions about the study:")
        doc.add_paragraph(f"Principal Investigator: {context.get('pi_name', '[PI NAME]')}")
        doc.add_paragraph(f"Phone: {context.get('pi_phone', '[PHONE]')}")
        doc.add_paragraph("")
        doc.add_paragraph("For questions about your rights as a research subject:")
        doc.add_paragraph(f"IRB: {context.get('irb_name', '[IRB NAME]')}")
        doc.add_paragraph(f"Phone: {context.get('irb_phone', '[IRB PHONE]')}")

        # Signature section
        doc.add_page_break()
        doc.add_heading("SIGNATURE PAGE", level=1)
        doc.add_paragraph("")
        doc.add_paragraph("I have read this consent form and have had the opportunity to ask questions.")
        doc.add_paragraph("I voluntarily agree to take part in this research study.")
        doc.add_paragraph("")
        doc.add_paragraph("_" * 50)
        doc.add_paragraph("Participant Name (printed)")
        doc.add_paragraph("")
        doc.add_paragraph("_" * 50 + "    Date: _____________")
        doc.add_paragraph("Participant Signature")
        doc.add_paragraph("")
        doc.add_paragraph("_" * 50 + "    Date: _____________")
        doc.add_paragraph("Person Obtaining Consent")
