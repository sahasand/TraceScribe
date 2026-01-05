"""DMP (Data Management Plan) generator."""

import json
import logging
from dataclasses import dataclass, field
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .base import BaseDocumentGenerator
from app.modules.ai.prompts.dmp_generation import DMP_GENERATION_PROMPT

logger = logging.getLogger(__name__)


@dataclass
class DMPContent:
    """DMP content structure."""
    protocol_number: str = ""
    study_title: str = ""
    sponsor: str = ""
    edc_system: str = "Medidata Rave"
    meddra_version: str = "26.1"
    whodrug_version: str = "March 2024"
    planned_enrollment: int = 0
    planned_sites: int = 0
    purpose_and_scope: str = ""
    study_information: str = ""
    roles_and_responsibilities: str = ""
    database_design: str = ""
    data_entry: str = ""
    data_validation: str = ""
    medical_coding: str = ""
    query_management: str = ""
    sae_reconciliation: str = ""
    external_data: str = ""
    database_lock: str = ""
    data_transfer: str = ""
    archiving: str = ""
    visits: list = field(default_factory=list)
    procedures: list = field(default_factory=list)


class DMPGenerator(BaseDocumentGenerator[DMPContent]):
    """Generator for Data Management Plans."""

    template_name = "standard_dmp.docx"
    document_type = "dmp"
    requires_polish = False

    async def extract_for_document(self, protocol_data: dict) -> DMPContent:
        """Extract DMP-relevant data from protocol."""
        metadata = protocol_data.get("metadata", {})
        design = protocol_data.get("design", {})
        visits = protocol_data.get("visits", [])
        procedures = protocol_data.get("procedures", [])

        # Generate DMP content using Gemini
        if self.gemini:
            generated = await self._generate_dmp_content(protocol_data)
        else:
            generated = self._generate_fallback_content(protocol_data)

        return DMPContent(
            protocol_number=metadata.get("protocol_number", ""),
            study_title=metadata.get("title", ""),
            sponsor=metadata.get("sponsor", ""),
            planned_enrollment=design.get("planned_enrollment", 0) or 0,
            purpose_and_scope=generated.get("purpose_and_scope", ""),
            study_information=generated.get("study_information", ""),
            roles_and_responsibilities=generated.get("roles_and_responsibilities", ""),
            database_design=generated.get("database_design", ""),
            data_entry=generated.get("data_entry", ""),
            data_validation=generated.get("data_validation", ""),
            medical_coding=generated.get("medical_coding", ""),
            query_management=generated.get("query_management", ""),
            sae_reconciliation=generated.get("sae_reconciliation", ""),
            external_data=generated.get("external_data", ""),
            database_lock=generated.get("database_lock", ""),
            data_transfer=generated.get("data_transfer", ""),
            archiving=generated.get("archiving", ""),
            visits=visits,
            procedures=procedures,
        )

    async def _generate_dmp_content(self, protocol_data: dict) -> dict:
        """Generate DMP content using Gemini."""
        prompt = DMP_GENERATION_PROMPT.format(
            protocol_data=json.dumps(protocol_data, indent=2)
        )

        try:
            response = await self.gemini.generate(prompt, temperature=0.2)

            json_str = response.strip()
            if json_str.startswith("```"):
                json_str = json_str.split("```")[1]
                if json_str.startswith("json"):
                    json_str = json_str[4:]
            json_str = json_str.strip()

            return json.loads(json_str)
        except Exception as e:
            logger.error(f"Failed to generate DMP content: {e}")
            return self._generate_fallback_content(protocol_data)

    def _generate_fallback_content(self, protocol_data: dict) -> dict:
        """Generate fallback DMP content without AI."""
        metadata = protocol_data.get("metadata", {})
        design = protocol_data.get("design", {})

        return {
            "purpose_and_scope": f"This Data Management Plan (DMP) describes the data management activities for protocol {metadata.get('protocol_number', 'TBD')}. The DMP defines the processes and procedures for data collection, entry, validation, query management, coding, and database lock.",
            "study_information": f"This is a {design.get('study_type', 'clinical')} study evaluating {metadata.get('indication', 'the investigational product')}. The study design is {design.get('design', 'as specified in the protocol')}.",
            "roles_and_responsibilities": "The Data Manager is responsible for database design, edit checks, and query management. The Medical Coder handles AE/MedHx coding and dictionary management. The Biostatistician provides SAP oversight and TLF review.",
            "database_design": "The clinical database will be designed in the EDC system based on the protocol visit schedule and assessments. Case Report Forms (CRFs) will be developed for each study visit.",
            "data_entry": "Data entry will be performed by trained site personnel directly into the EDC system. Double data entry is not required due to built-in edit checks.",
            "data_validation": "Edit checks will be programmed to identify data discrepancies, out-of-range values, and protocol deviations. Queries will be generated automatically for review by site staff.",
            "medical_coding": "Adverse events will be coded using MedDRA. Concomitant medications will be coded using WHODrug.",
            "query_management": "Data queries will be managed through the EDC system. Sites are expected to respond to queries within 5 business days.",
            "sae_reconciliation": "SAE reconciliation between the clinical database and safety database will be performed monthly.",
            "external_data": "External data (e.g., central lab, ECG) will be transferred electronically and reconciled with eCRF data.",
            "database_lock": "Database lock will occur after all data has been entered, queries resolved, and medical coding completed. The database lock meeting will include Data Management, Biostatistics, and Clinical Operations.",
            "data_transfer": "Data transfers to the Sponsor will be performed using secure file transfer protocols. Transfer specifications will be documented in a separate Data Transfer Agreement.",
            "archiving": "All study data and documentation will be archived according to regulatory requirements and sponsor SOPs.",
        }

    def build_template_context(self, extracted: DMPContent) -> dict:
        """Build context for DMP template."""
        return {
            "title": "DATA MANAGEMENT PLAN",
            "protocol_number": extracted.protocol_number,
            "study_title": extracted.study_title,
            "sponsor": extracted.sponsor,
            "edc_system": extracted.edc_system,
            "meddra_version": extracted.meddra_version,
            "whodrug_version": extracted.whodrug_version,
            "planned_enrollment": extracted.planned_enrollment,
            "planned_sites": extracted.planned_sites,
            "purpose_and_scope": extracted.purpose_and_scope,
            "study_information": extracted.study_information,
            "roles_and_responsibilities": extracted.roles_and_responsibilities,
            "database_design": extracted.database_design,
            "data_entry": extracted.data_entry,
            "data_validation": extracted.data_validation,
            "medical_coding": extracted.medical_coding,
            "query_management": extracted.query_management,
            "sae_reconciliation": extracted.sae_reconciliation,
            "external_data": extracted.external_data,
            "database_lock": extracted.database_lock,
            "data_transfer": extracted.data_transfer,
            "archiving": extracted.archiving,
            "visits": extracted.visits,
            "procedures": extracted.procedures,
        }

    async def _add_content_to_document(self, doc: Document, context: dict) -> None:
        """Add DMP content to document with 4-level numbering."""
        # Title page
        title = doc.add_heading("DATA MANAGEMENT PLAN", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")
        doc.add_paragraph(f"Protocol Number: {context.get('protocol_number', '')}")
        doc.add_paragraph(f"Study Title: {context.get('study_title', '')}")
        doc.add_paragraph(f"Sponsor: {context.get('sponsor', '')}")
        doc.add_paragraph("")
        doc.add_paragraph(f"EDC System: {context.get('edc_system', 'Medidata Rave')}")
        doc.add_paragraph("")

        # Version history table
        doc.add_heading("Version History", level=2)
        version_table = doc.add_table(rows=2, cols=4)
        version_table.style = "Table Grid"
        headers = ["Version", "Date", "Author", "Description"]
        for i, header in enumerate(headers):
            version_table.rows[0].cells[i].text = header
        version_table.rows[1].cells[0].text = "1.0"
        version_table.rows[1].cells[1].text = "[Date]"
        version_table.rows[1].cells[2].text = "[Author]"
        version_table.rows[1].cells[3].text = "Initial version"

        doc.add_page_break()

        # Main sections with 4-level numbering
        sections = [
            ("1", "PURPOSE AND SCOPE", context.get("purpose_and_scope", "")),
            ("2", "STUDY INFORMATION", context.get("study_information", "")),
            ("3", "ROLES AND RESPONSIBILITIES", context.get("roles_and_responsibilities", "")),
            ("4", "DATABASE DESIGN", context.get("database_design", "")),
            ("5", "DATA ENTRY", context.get("data_entry", "")),
            ("6", "DATA VALIDATION (EDIT CHECKS)", context.get("data_validation", "")),
            ("7", "MEDICAL CODING", context.get("medical_coding", "")),
            ("8", "DATA REVIEW AND QUERY MANAGEMENT", context.get("query_management", "")),
            ("9", "SAE RECONCILIATION", context.get("sae_reconciliation", "")),
            ("10", "EXTERNAL DATA MANAGEMENT", context.get("external_data", "")),
            ("11", "DATABASE LOCK PROCEDURES", context.get("database_lock", "")),
            ("12", "DATA TRANSFER", context.get("data_transfer", "")),
            ("13", "ARCHIVING", context.get("archiving", "")),
        ]

        for num, heading, content in sections:
            doc.add_heading(f"{num}. {heading}", level=1)
            doc.add_paragraph(content)
            doc.add_paragraph("")

        # Add roles and responsibilities table
        doc.add_heading("3.1 Roles Table", level=2)
        roles_table = doc.add_table(rows=4, cols=2)
        roles_table.style = "Table Grid"
        roles_table.rows[0].cells[0].text = "Role"
        roles_table.rows[0].cells[1].text = "Responsibility"
        roles_table.rows[1].cells[0].text = "Data Manager"
        roles_table.rows[1].cells[1].text = "Database design, edit checks, query management"
        roles_table.rows[2].cells[0].text = "Medical Coder"
        roles_table.rows[2].cells[1].text = "AE/MedHx coding, dictionary management"
        roles_table.rows[3].cells[0].text = "Biostatistician"
        roles_table.rows[3].cells[1].text = "SAP oversight, TLF review"

        doc.add_paragraph("")

        # Dictionary versions table
        doc.add_heading("7.1 Dictionary Versions", level=2)
        dict_table = doc.add_table(rows=3, cols=2)
        dict_table.style = "Table Grid"
        dict_table.rows[0].cells[0].text = "Dictionary"
        dict_table.rows[0].cells[1].text = "Version"
        dict_table.rows[1].cells[0].text = "MedDRA"
        dict_table.rows[1].cells[1].text = context.get("meddra_version", "26.1")
        dict_table.rows[2].cells[0].text = "WHODrug"
        dict_table.rows[2].cells[1].text = context.get("whodrug_version", "March 2024")

        # Visit schedule if available
        if context.get("visits"):
            doc.add_page_break()
            doc.add_heading("14. APPENDIX A: VISIT SCHEDULE", level=1)

            visits = context["visits"]
            if visits:
                visit_table = doc.add_table(rows=len(visits) + 1, cols=3)
                visit_table.style = "Table Grid"
                visit_table.rows[0].cells[0].text = "Visit"
                visit_table.rows[0].cells[1].text = "Timing"
                visit_table.rows[0].cells[2].text = "Procedures"

                for i, visit in enumerate(visits, 1):
                    visit_table.rows[i].cells[0].text = visit.get("name", "")
                    visit_table.rows[i].cells[1].text = visit.get("timing", "")
                    visit_table.rows[i].cells[2].text = ", ".join(visit.get("procedures", []))
