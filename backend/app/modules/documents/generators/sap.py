"""SAP (Statistical Analysis Plan) generator."""

import json
import logging
from dataclasses import dataclass, field
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .base import BaseDocumentGenerator
from app.modules.ai.prompts.sap_generation import SAP_GENERATION_PROMPT

logger = logging.getLogger(__name__)


@dataclass
class SAPContent:
    """SAP content structure."""
    protocol_number: str = ""
    study_title: str = ""
    sponsor: str = ""
    phase: str = ""
    study_design: str = ""
    planned_enrollment: int = 0
    randomization_ratio: str = ""
    introduction: str = ""
    primary_objective: str = ""
    primary_endpoints: list = field(default_factory=list)  # VERBATIM
    secondary_objectives: str = ""
    secondary_endpoints: list = field(default_factory=list)  # VERBATIM
    study_design_summary: str = ""
    analysis_populations: dict = field(default_factory=dict)
    statistical_methods: dict = field(default_factory=dict)
    sample_size: str = ""
    missing_data: str = ""
    interim_analysis: str = ""
    tlf_shells: list = field(default_factory=list)


class SAPGenerator(BaseDocumentGenerator[SAPContent]):
    """Generator for Statistical Analysis Plans."""

    template_name = "standard_sap.docx"
    document_type = "sap"
    requires_polish = True

    async def extract_for_document(self, protocol_data: dict) -> SAPContent:
        """Extract SAP-relevant data from protocol."""
        metadata = protocol_data.get("metadata", {})
        design = protocol_data.get("design", {})
        endpoints = protocol_data.get("endpoints", {})

        # Generate SAP content using Gemini
        if self.gemini:
            generated = await self._generate_sap_content(protocol_data)
        else:
            generated = self._generate_fallback_content(protocol_data)

        return SAPContent(
            protocol_number=metadata.get("protocol_number", ""),
            study_title=metadata.get("title", ""),
            sponsor=metadata.get("sponsor", ""),
            phase=metadata.get("phase", ""),
            study_design=design.get("design", ""),
            planned_enrollment=design.get("planned_enrollment", 0) or 0,
            randomization_ratio=design.get("randomization_ratio", ""),
            introduction=generated.get("introduction", ""),
            primary_objective=generated.get("objectives_and_endpoints", {}).get("primary_objective", ""),
            # VERBATIM endpoints from protocol
            primary_endpoints=endpoints.get("primary", []),
            secondary_objectives=generated.get("objectives_and_endpoints", {}).get("secondary_objectives", ""),
            secondary_endpoints=endpoints.get("secondary", []),
            study_design_summary=generated.get("study_design_summary", ""),
            analysis_populations=generated.get("analysis_populations", {}),
            statistical_methods=generated.get("statistical_methods", {}),
            sample_size=generated.get("sample_size", ""),
            missing_data=generated.get("missing_data", ""),
            interim_analysis=generated.get("interim_analysis", ""),
            tlf_shells=generated.get("tlf_shells", []),
        )

    async def _generate_sap_content(self, protocol_data: dict) -> dict:
        """Generate SAP content using Gemini."""
        prompt = SAP_GENERATION_PROMPT.format(
            protocol_data=json.dumps(protocol_data, indent=2)
        )

        try:
            response = await self.gemini.generate(prompt, temperature=0.2)

            json_str = response.strip()
            if json_str.startswith("```"):
                json_str = json_str.split("```")[1]
                if json_str.startswith("json"):
                    json_str = json_str[4:]
            json_str = json_str.strip()

            return json.loads(json_str)
        except Exception as e:
            logger.error(f"Failed to generate SAP content: {e}")
            return self._generate_fallback_content(protocol_data)

    def _generate_fallback_content(self, protocol_data: dict) -> dict:
        """Generate fallback SAP content without AI."""
        metadata = protocol_data.get("metadata", {})
        design = protocol_data.get("design", {})
        endpoints = protocol_data.get("endpoints", {})

        return {
            "introduction": f"This Statistical Analysis Plan (SAP) describes the planned statistical analyses for protocol {metadata.get('protocol_number', 'TBD')}. This SAP should be read in conjunction with the study protocol and has been developed in accordance with ICH E9 Statistical Principles for Clinical Trials.",
            "objectives_and_endpoints": {
                "primary_objective": f"To evaluate the efficacy and safety of the investigational product in patients with {metadata.get('indication', 'the target indication')}.",
                "primary_endpoints": endpoints.get("primary", []),
                "secondary_objectives": "To evaluate secondary efficacy measures and characterize the safety profile.",
                "secondary_endpoints": endpoints.get("secondary", []),
            },
            "study_design_summary": f"This is a {design.get('study_type', 'clinical')} study with {design.get('blinding', 'the specified')} design. Approximately {design.get('planned_enrollment', 'N')} subjects will be enrolled.",
            "analysis_populations": {
                "itt": "Intent-to-Treat (ITT): All randomized subjects",
                "mitt": "Modified ITT (mITT): All randomized subjects who received at least one dose of study drug",
                "per_protocol": "Per-Protocol (PP): All subjects who completed the study without major protocol deviations",
                "safety": "Safety Population: All subjects who received at least one dose of study drug",
            },
            "statistical_methods": {
                "general": "All statistical analyses will be performed using SAS Version 9.4. A two-sided significance level of 0.05 will be used unless otherwise specified.",
                "primary_analysis": "The primary efficacy analysis will be based on the ITT population. The primary endpoint will be analyzed using appropriate statistical methods based on the endpoint type.",
                "secondary_analysis": "Secondary endpoints will be analyzed using appropriate statistical methods. No adjustments for multiplicity will be made for secondary endpoints.",
                "safety_analysis": "Safety analyses will be conducted on the Safety Population. Adverse events will be summarized by System Organ Class and Preferred Term.",
            },
            "sample_size": f"The planned sample size is {design.get('planned_enrollment', 'N')} subjects. Sample size calculations are provided in the protocol.",
            "missing_data": "Missing data will be handled using appropriate methods. The primary analysis will use observed data. Sensitivity analyses may include multiple imputation or last observation carried forward.",
            "interim_analysis": "No formal interim analysis is planned unless specified in the protocol.",
            "tlf_shells": [
                "Table 14.1.1: Subject Disposition",
                "Table 14.1.2: Demographics and Baseline Characteristics",
                "Table 14.2.1: Primary Efficacy Analysis",
                "Table 14.3.1: Overall Summary of Adverse Events",
                "Table 14.3.2: Adverse Events by System Organ Class and Preferred Term",
                "Listing 16.2.1: Subject Demographics",
                "Listing 16.2.2: Adverse Events",
                "Figure 14.2.1: Primary Endpoint Over Time",
            ],
        }

    def build_template_context(self, extracted: SAPContent) -> dict:
        """Build context for SAP template."""
        return {
            "title": "STATISTICAL ANALYSIS PLAN",
            "protocol_number": extracted.protocol_number,
            "study_title": extracted.study_title,
            "sponsor": extracted.sponsor,
            "phase": extracted.phase,
            "study_design": extracted.study_design,
            "planned_enrollment": extracted.planned_enrollment,
            "randomization_ratio": extracted.randomization_ratio,
            "introduction": extracted.introduction,
            "primary_objective": extracted.primary_objective,
            "primary_endpoints": extracted.primary_endpoints,
            "secondary_objectives": extracted.secondary_objectives,
            "secondary_endpoints": extracted.secondary_endpoints,
            "study_design_summary": extracted.study_design_summary,
            "analysis_populations": extracted.analysis_populations,
            "statistical_methods": extracted.statistical_methods,
            "sample_size": extracted.sample_size,
            "missing_data": extracted.missing_data,
            "interim_analysis": extracted.interim_analysis,
            "tlf_shells": extracted.tlf_shells,
        }

    async def _add_content_to_document(self, doc: Document, context: dict) -> None:
        """Add SAP content to document."""
        # Title page
        title = doc.add_heading("STATISTICAL ANALYSIS PLAN", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")
        doc.add_paragraph(f"Protocol Number: {context.get('protocol_number', '')}")
        doc.add_paragraph(f"Study Title: {context.get('study_title', '')}")
        doc.add_paragraph(f"Sponsor: {context.get('sponsor', '')}")
        doc.add_paragraph(f"Phase: {context.get('phase', '')}")
        doc.add_paragraph("")

        # Version history
        doc.add_heading("Version History", level=2)
        version_table = doc.add_table(rows=2, cols=4)
        version_table.style = "Table Grid"
        headers = ["Version", "Date", "Author", "Description"]
        for i, header in enumerate(headers):
            version_table.rows[0].cells[i].text = header
        version_table.rows[1].cells[0].text = "1.0"
        version_table.rows[1].cells[1].text = "[Date]"
        version_table.rows[1].cells[2].text = "[Biostatistician]"
        version_table.rows[1].cells[3].text = "Initial version"

        doc.add_page_break()

        # Table of Contents placeholder
        doc.add_heading("TABLE OF CONTENTS", level=1)
        doc.add_paragraph("[Table of Contents will be generated in final document]")
        doc.add_page_break()

        # 1. Introduction
        doc.add_heading("1. INTRODUCTION", level=1)
        doc.add_paragraph(context.get("introduction", ""))
        doc.add_paragraph("")

        # 2. Study Objectives and Endpoints
        doc.add_heading("2. STUDY OBJECTIVES AND ENDPOINTS", level=1)

        doc.add_heading("2.1 Primary Objective", level=2)
        doc.add_paragraph(context.get("primary_objective", ""))

        doc.add_heading("2.2 Primary Endpoints", level=2)
        doc.add_paragraph("The primary endpoints are (verbatim from protocol):")
        for i, endpoint in enumerate(context.get("primary_endpoints", []), 1):
            doc.add_paragraph(f"{i}. {endpoint}")

        doc.add_heading("2.3 Secondary Objectives", level=2)
        doc.add_paragraph(context.get("secondary_objectives", ""))

        doc.add_heading("2.4 Secondary Endpoints", level=2)
        doc.add_paragraph("The secondary endpoints are (verbatim from protocol):")
        for i, endpoint in enumerate(context.get("secondary_endpoints", []), 1):
            doc.add_paragraph(f"{i}. {endpoint}")

        # 3. Study Design
        doc.add_heading("3. STUDY DESIGN", level=1)
        doc.add_paragraph(context.get("study_design_summary", ""))
        doc.add_paragraph("")
        doc.add_paragraph(f"Planned Enrollment: {context.get('planned_enrollment', 'N/A')}")
        doc.add_paragraph(f"Randomization Ratio: {context.get('randomization_ratio', 'N/A')}")

        # 4. Analysis Populations
        doc.add_heading("4. ANALYSIS POPULATIONS", level=1)
        populations = context.get("analysis_populations", {})

        pop_table = doc.add_table(rows=5, cols=2)
        pop_table.style = "Table Grid"
        pop_table.rows[0].cells[0].text = "Population"
        pop_table.rows[0].cells[1].text = "Definition"
        pop_table.rows[1].cells[0].text = "ITT"
        pop_table.rows[1].cells[1].text = populations.get("itt", "All randomized subjects")
        pop_table.rows[2].cells[0].text = "mITT"
        pop_table.rows[2].cells[1].text = populations.get("mitt", "Randomized + received ≥1 dose")
        pop_table.rows[3].cells[0].text = "Per-Protocol"
        pop_table.rows[3].cells[1].text = populations.get("per_protocol", "Completed without major deviations")
        pop_table.rows[4].cells[0].text = "Safety"
        pop_table.rows[4].cells[1].text = populations.get("safety", "Received ≥1 dose of study drug")

        # 5. Statistical Methods
        doc.add_heading("5. STATISTICAL METHODS", level=1)
        methods = context.get("statistical_methods", {})

        doc.add_heading("5.1 General Considerations", level=2)
        doc.add_paragraph(methods.get("general", ""))

        doc.add_heading("5.2 Primary Endpoint Analysis", level=2)
        doc.add_paragraph(methods.get("primary_analysis", ""))

        doc.add_heading("5.3 Secondary Endpoint Analysis", level=2)
        doc.add_paragraph(methods.get("secondary_analysis", ""))

        doc.add_heading("5.4 Safety Analysis", level=2)
        doc.add_paragraph(methods.get("safety_analysis", ""))

        # 6. Sample Size
        doc.add_heading("6. SAMPLE SIZE", level=1)
        doc.add_paragraph(context.get("sample_size", ""))

        # 7. Handling Missing Data
        doc.add_heading("7. HANDLING OF MISSING DATA", level=1)
        doc.add_paragraph(context.get("missing_data", ""))

        # 8. Interim Analysis
        doc.add_heading("8. INTERIM ANALYSIS", level=1)
        doc.add_paragraph(context.get("interim_analysis", ""))

        # 9. Tables, Listings, and Figures
        doc.add_heading("9. TABLES, LISTINGS, AND FIGURES SHELL", level=1)
        doc.add_paragraph("The following TLFs will be produced:")
        doc.add_paragraph("")

        for tlf in context.get("tlf_shells", []):
            doc.add_paragraph(f"• {tlf}", style="List Bullet")

        # Signature page
        doc.add_page_break()
        doc.add_heading("SIGNATURE PAGE", level=1)
        doc.add_paragraph("")
        doc.add_paragraph("This Statistical Analysis Plan has been reviewed and approved.")
        doc.add_paragraph("")
        doc.add_paragraph("_" * 50 + "    Date: _____________")
        doc.add_paragraph("Lead Biostatistician")
        doc.add_paragraph("")
        doc.add_paragraph("_" * 50 + "    Date: _____________")
        doc.add_paragraph("Medical Monitor")
        doc.add_paragraph("")
        doc.add_paragraph("_" * 50 + "    Date: _____________")
        doc.add_paragraph("Sponsor Representative")
