"""List builder for the Unified Document Engine.

This module provides the ListBuilder class for rendering bullet and numbered
lists from UIF (Unified Intermediate Format) list data to python-docx documents.

Supports:
- Bullet lists with proper styling
- Numbered lists (decimal, roman numerals, letters)
- Multi-level/nested lists up to 9 levels deep
- Inline formatting within list items (bold, italic, underline)
- Proper hanging indents for list items
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from typing import Any, Optional, Union

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap

logger = logging.getLogger(__name__)


# Constants for list formatting
INDENT_PER_LEVEL = Inches(0.25)  # Indentation per nesting level
HANGING_INDENT = Inches(0.25)   # Hanging indent for list items
MAX_NESTING_LEVEL = 9           # Maximum supported nesting depth


@dataclass
class InlineFormatting:
    """Inline formatting specification for text runs within list items.

    Attributes:
        bold: Apply bold formatting to text
        italic: Apply italic formatting to text
        underline: Apply underline formatting to text
        strike: Apply strikethrough formatting
        superscript: Apply superscript formatting
        subscript: Apply subscript formatting
        font_name: Font family name
        font_size: Font size in points
    """
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    strike: Optional[bool] = None
    superscript: Optional[bool] = None
    subscript: Optional[bool] = None
    font_name: Optional[str] = None
    font_size: Optional[float] = None


@dataclass
class TextRun:
    """A run of text with optional formatting.

    Attributes:
        text: The text content
        formatting: Optional inline formatting to apply
    """
    text: str
    formatting: Optional[InlineFormatting] = None


@dataclass
class ListItem:
    """A single item within a list.

    Attributes:
        content: The text content (can be plain string or list of TextRuns)
        level: Nesting level (0 = top level, 1 = first nested level, etc.)
        formatting: Default formatting for the entire item
        runs: List of TextRuns for mixed formatting within the item
    """
    content: Union[str, list[TextRun]]
    level: int = 0
    formatting: Optional[InlineFormatting] = None
    runs: list[TextRun] = field(default_factory=list)

    def get_text_runs(self) -> list[TextRun]:
        """Get content as list of TextRuns for rendering.

        Returns:
            List of TextRun objects representing the content
        """
        if self.runs:
            return self.runs
        if isinstance(self.content, str):
            return [TextRun(text=self.content, formatting=self.formatting)]
        return self.content


@dataclass
class ListBlock:
    """A complete list block in UIF format.

    Attributes:
        type: Block type ("bullet_list" or "numbered_list")
        list_style: Style variant for the list
        items: List items (can be strings, ListItem objects, or dicts)
        start_number: Starting number for numbered lists (default 1)
    """
    type: str  # "bullet_list" or "numbered_list"
    list_style: str = "bullet"  # "bullet", "decimal", "roman_upper", "roman_lower", "letter_upper", "letter_lower"
    items: list[Union[str, ListItem, dict]] = field(default_factory=list)
    start_number: int = 1


class ListBuilder:
    """Builder class for rendering lists to Word documents.

    This builder takes UIF list data and renders it to python-docx Document
    objects with proper formatting, indentation, and numbering.

    Example:
        >>> from docx import Document
        >>> doc = Document()
        >>> builder = ListBuilder(doc)
        >>> builder.add_bullet_list(["Item 1", "Item 2", "Item 3"])
        >>> builder.add_numbered_list(["First", "Second"], style="roman_lower")
    """

    # Mapping of list styles to their numbering format identifiers
    NUMBERING_FORMATS = {
        "decimal": "decimal",
        "roman_upper": "upperRoman",
        "roman_lower": "lowerRoman",
        "letter_upper": "upperLetter",
        "letter_lower": "lowerLetter",
        "bullet": "bullet",
    }

    # Bullet characters for different levels
    BULLET_CHARS = [
        "\u2022",  # Level 0: Bullet (filled circle)
        "\u25E6",  # Level 1: White bullet (hollow circle)
        "\u25AA",  # Level 2: Black small square
        "\u2022",  # Level 3: Bullet
        "\u25E6",  # Level 4: White bullet
        "\u25AA",  # Level 5: Black small square
        "\u2022",  # Level 6: Bullet
        "\u25E6",  # Level 7: White bullet
        "\u25AA",  # Level 8: Black small square
    ]

    def __init__(self, doc: Document):
        """Initialize the ListBuilder.

        Args:
            doc: The python-docx Document to add lists to
        """
        self.doc = doc
        self._numbering_instance_counter = 0
        self._ensure_list_styles()

    def add_list(self, block: Union[dict, ListBlock]) -> None:
        """Add a list from UIF block data.

        Args:
            block: UIF ListBlock data (dict or ListBlock object)

        Raises:
            ValueError: If block type is not a valid list type
        """
        # Convert dict to ListBlock if needed
        if isinstance(block, dict):
            block_type = block.get("type", "bullet_list")
            list_style = block.get("list_style", "bullet")
            items = block.get("items", [])
            start_number = block.get("start_number", 1)
        else:
            block_type = block.type
            list_style = block.list_style
            items = block.items
            start_number = block.start_number

        if block_type == "bullet_list":
            self.add_bullet_list(items)
        elif block_type == "numbered_list":
            self.add_numbered_list(items, style=list_style, start_number=start_number)
        else:
            raise ValueError(f"Unknown list type: {block_type}")

    def add_bullet_list(self, items: list) -> None:
        """Add a bullet list to the document.

        Args:
            items: List of items (strings, ListItem objects, or dicts)
        """
        for item in items:
            parsed_item = self._parse_item(item)
            style_name = self._get_style_name("bullet", parsed_item.level)
            self._add_list_item(parsed_item, style_name, is_bullet=True)

    def add_numbered_list(
        self,
        items: list,
        style: str = "decimal",
        start_number: int = 1
    ) -> None:
        """Add a numbered list with specified style.

        Args:
            items: List of items (strings, ListItem objects, or dicts)
            style: Numbering style - one of:
                - "decimal": 1, 2, 3...
                - "roman_upper": I, II, III...
                - "roman_lower": i, ii, iii...
                - "letter_upper": A, B, C...
                - "letter_lower": a, b, c...
            start_number: Starting number (default 1)
        """
        if style not in self.NUMBERING_FORMATS:
            logger.warning(f"Unknown numbering style '{style}', defaulting to 'decimal'")
            style = "decimal"

        # Create a new numbering instance for this list
        num_id = self._create_numbering_definition(style, start_number)

        for item in items:
            parsed_item = self._parse_item(item)
            self._add_numbered_item(parsed_item, num_id, style)

    def _parse_item(self, item: Union[str, ListItem, dict]) -> ListItem:
        """Parse an item into a ListItem object.

        Args:
            item: Item data in various formats

        Returns:
            Parsed ListItem object
        """
        if isinstance(item, ListItem):
            return item
        elif isinstance(item, str):
            return ListItem(content=item, level=0)
        elif isinstance(item, dict):
            content = item.get("content", "")
            level = item.get("level", 0)
            formatting = None
            runs = []

            # Parse formatting
            if "formatting" in item and item["formatting"]:
                fmt_data = item["formatting"]
                if isinstance(fmt_data, InlineFormatting):
                    formatting = fmt_data
                elif isinstance(fmt_data, dict):
                    formatting = InlineFormatting(**fmt_data)

            # Parse runs
            if "runs" in item:
                for run_data in item["runs"]:
                    if isinstance(run_data, TextRun):
                        runs.append(run_data)
                    elif isinstance(run_data, dict):
                        run_formatting = None
                        if "formatting" in run_data and run_data["formatting"]:
                            run_fmt = run_data["formatting"]
                            if isinstance(run_fmt, InlineFormatting):
                                run_formatting = run_fmt
                            elif isinstance(run_fmt, dict):
                                run_formatting = InlineFormatting(**run_fmt)
                        runs.append(TextRun(
                            text=run_data.get("text", ""),
                            formatting=run_formatting
                        ))

            return ListItem(
                content=content,
                level=min(level, MAX_NESTING_LEVEL - 1),
                formatting=formatting,
                runs=runs
            )
        else:
            # Fallback: convert to string
            return ListItem(content=str(item), level=0)

    def _add_list_item(
        self,
        item: ListItem,
        style_name: str,
        is_bullet: bool = True
    ) -> None:
        """Add a single list item with bullet formatting.

        Args:
            item: The ListItem to add
            style_name: Word style name to use
            is_bullet: Whether this is a bullet item
        """
        # Create paragraph
        para = self.doc.add_paragraph()

        # Apply style if it exists
        try:
            para.style = style_name
        except KeyError:
            # Style doesn't exist, apply manual formatting
            pass

        # Set paragraph formatting for proper list appearance
        pf = para.paragraph_format
        level = item.level

        # Calculate indentation
        left_indent = INDENT_PER_LEVEL * (level + 1)
        pf.left_indent = left_indent
        pf.first_line_indent = -HANGING_INDENT

        # Set spacing
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Add bullet character manually for bullet lists
        if is_bullet:
            bullet_char = self.BULLET_CHARS[level % len(self.BULLET_CHARS)]
            bullet_run = para.add_run(f"{bullet_char}\t")
            bullet_run.font.size = Pt(11)

        # Add content with formatting
        text_runs = item.get_text_runs()
        for text_run in text_runs:
            run = para.add_run(text_run.text)
            self._apply_run_formatting(run, text_run.formatting)

    def _add_numbered_item(
        self,
        item: ListItem,
        num_id: int,
        style: str
    ) -> None:
        """Add a single numbered list item.

        Args:
            item: The ListItem to add
            num_id: Numbering definition ID
            style: Numbering style
        """
        para = self.doc.add_paragraph()

        # Apply numbering through XML
        self._apply_numbering_to_paragraph(para, num_id, item.level)

        # Set paragraph formatting
        pf = para.paragraph_format
        level = item.level

        # Calculate indentation
        left_indent = INDENT_PER_LEVEL * (level + 1)
        pf.left_indent = left_indent
        pf.first_line_indent = -HANGING_INDENT

        # Set spacing
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Add content with formatting
        text_runs = item.get_text_runs()
        for text_run in text_runs:
            run = para.add_run(text_run.text)
            self._apply_run_formatting(run, text_run.formatting)

    def _apply_run_formatting(
        self,
        run,
        formatting: Optional[InlineFormatting]
    ) -> None:
        """Apply inline formatting to a run.

        Args:
            run: The docx Run object to format
            formatting: Formatting specification to apply
        """
        if formatting is None:
            # Apply default font size
            run.font.size = Pt(11)
            return

        # Apply each formatting option if specified
        if formatting.bold is not None:
            run.font.bold = formatting.bold
        if formatting.italic is not None:
            run.font.italic = formatting.italic
        if formatting.underline is not None:
            run.font.underline = formatting.underline
        if formatting.strike is not None:
            run.font.strike = formatting.strike
        if formatting.superscript is not None:
            run.font.superscript = formatting.superscript
        if formatting.subscript is not None:
            run.font.subscript = formatting.subscript
        if formatting.font_name is not None:
            run.font.name = formatting.font_name
        if formatting.font_size is not None:
            run.font.size = Pt(formatting.font_size)
        else:
            run.font.size = Pt(11)

    def _ensure_list_styles(self) -> None:
        """Ensure required list styles exist in the document.

        Creates custom list styles if they don't exist in the document's
        style definitions.
        """
        styles = self.doc.styles

        # Define styles we need for bullet lists at different levels
        bullet_styles = [
            f"ListBulletLevel{i}" for i in range(MAX_NESTING_LEVEL)
        ]

        # Define styles for numbered lists
        number_styles = [
            f"ListNumberLevel{i}" for i in range(MAX_NESTING_LEVEL)
        ]

        all_styles = bullet_styles + number_styles

        for style_name in all_styles:
            try:
                # Check if style exists
                _ = styles[style_name]
            except KeyError:
                # Style doesn't exist, create it
                try:
                    # Create based on Normal style
                    style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                    style.base_style = styles['Normal']
                    style.font.size = Pt(11)

                    # Extract level from style name
                    level = int(style_name[-1])

                    # Set paragraph formatting
                    pf = style.paragraph_format
                    pf.left_indent = INDENT_PER_LEVEL * (level + 1)
                    pf.first_line_indent = -HANGING_INDENT
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(3)

                    logger.debug(f"Created list style: {style_name}")
                except Exception as e:
                    logger.debug(f"Could not create style {style_name}: {e}")

    def _get_style_name(self, list_type: str, level: int) -> str:
        """Get the Word style name for a list type and level.

        Args:
            list_type: "bullet" or numbering style name
            level: Nesting level (0-based)

        Returns:
            Style name string
        """
        level = min(level, MAX_NESTING_LEVEL - 1)

        if list_type == "bullet":
            return f"ListBulletLevel{level}"
        else:
            return f"ListNumberLevel{level}"

    def _create_numbering_definition(
        self,
        style: str,
        start_number: int = 1
    ) -> int:
        """Create a numbering definition for a numbered list.

        Args:
            style: Numbering style
            start_number: Starting number

        Returns:
            Numbering definition ID
        """
        # Get or create numbering part
        numbering_part = self._get_or_create_numbering_part()

        # Increment counter for unique IDs
        self._numbering_instance_counter += 1
        abstract_num_id = self._numbering_instance_counter
        num_id = self._numbering_instance_counter

        # Get the numbering format
        num_fmt = self.NUMBERING_FORMATS.get(style, "decimal")

        # Create abstract numbering definition
        abstract_num = self._create_abstract_numbering(
            abstract_num_id, num_fmt, start_number
        )

        # Add to numbering definitions
        numbering_elm = numbering_part.element

        # Insert abstract numbering before any num elements
        num_elements = numbering_elm.findall(qn('w:num'))
        if num_elements:
            num_elements[0].addprevious(abstract_num)
        else:
            numbering_elm.append(abstract_num)

        # Create num element that references the abstract numbering
        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(num_id))
        abstract_num_ref = OxmlElement('w:abstractNumId')
        abstract_num_ref.set(qn('w:val'), str(abstract_num_id))
        num.append(abstract_num_ref)
        numbering_elm.append(num)

        return num_id

    def _create_abstract_numbering(
        self,
        abstract_num_id: int,
        num_fmt: str,
        start_number: int
    ) -> OxmlElement:
        """Create an abstract numbering XML element.

        Args:
            abstract_num_id: Abstract numbering ID
            num_fmt: Numbering format string
            start_number: Starting number

        Returns:
            Abstract numbering OxmlElement
        """
        abstract_num = OxmlElement('w:abstractNum')
        abstract_num.set(qn('w:abstractNumId'), str(abstract_num_id))

        # Add multi-level type
        multi_level = OxmlElement('w:multiLevelType')
        multi_level.set(qn('w:val'), 'hybridMultilevel')
        abstract_num.append(multi_level)

        # Create level definitions for each nesting level
        for level in range(MAX_NESTING_LEVEL):
            lvl = self._create_level_definition(level, num_fmt, start_number)
            abstract_num.append(lvl)

        return abstract_num

    def _create_level_definition(
        self,
        level: int,
        num_fmt: str,
        start_number: int
    ) -> OxmlElement:
        """Create a level definition for numbering.

        Args:
            level: The level index (0-based)
            num_fmt: Numbering format
            start_number: Starting number for level 0

        Returns:
            Level OxmlElement
        """
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(level))

        # Start number (only meaningful for level 0)
        start = OxmlElement('w:start')
        start.set(qn('w:val'), str(start_number if level == 0 else 1))
        lvl.append(start)

        # Number format
        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), num_fmt)
        lvl.append(numFmt)

        # Level text (e.g., "1.", "a)", "i.")
        lvlText = OxmlElement('w:lvlText')
        text_pattern = self._get_level_text_pattern(level, num_fmt)
        lvlText.set(qn('w:val'), text_pattern)
        lvl.append(lvlText)

        # Level justification
        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)

        # Paragraph properties
        pPr = OxmlElement('w:pPr')

        # Indentation
        ind = OxmlElement('w:ind')
        left_indent = int(INDENT_PER_LEVEL.twips * (level + 1))
        hanging = int(HANGING_INDENT.twips)
        ind.set(qn('w:left'), str(left_indent))
        ind.set(qn('w:hanging'), str(hanging))
        pPr.append(ind)
        lvl.append(pPr)

        return lvl

    def _get_level_text_pattern(self, level: int, num_fmt: str) -> str:
        """Get the text pattern for a numbering level.

        Args:
            level: Level index
            num_fmt: Numbering format

        Returns:
            Text pattern string (e.g., "%1.", "(%2)")
        """
        level_ref = f"%{level + 1}"

        # Different patterns for different formats
        if num_fmt in ("decimal", "upperRoman", "lowerRoman"):
            return f"{level_ref}."
        elif num_fmt in ("upperLetter", "lowerLetter"):
            return f"{level_ref})"
        else:
            return f"{level_ref}."

    def _get_or_create_numbering_part(self):
        """Get or create the numbering part of the document.

        Returns:
            The numbering part object
        """
        # Try to get existing numbering part
        try:
            return self.doc.part.numbering_part
        except AttributeError:
            pass

        # Need to create one - add a dummy list item to trigger creation
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Access or create numbering definitions
        try:
            numbering_part = self.doc.part.numbering_part
        except:
            # Create a minimal numbering part
            from docx.parts.numbering import NumberingPart
            from docx.opc.constants import CONTENT_TYPE, RELATIONSHIP_TYPE

            # Create empty numbering XML
            numbering_xml = (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
                '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">\n'
                '</w:numbering>'
            )

            # Parse and create part
            from lxml import etree
            numbering_elm = etree.fromstring(numbering_xml.encode())

            numbering_part = NumberingPart.new()
            self.doc.part.relate_to(
                numbering_part,
                RELATIONSHIP_TYPE.NUMBERING
            )

        return self.doc.part.numbering_part

    def _apply_numbering_to_paragraph(
        self,
        paragraph,
        num_id: int,
        level: int
    ) -> None:
        """Apply numbering to a paragraph.

        Args:
            paragraph: The paragraph to apply numbering to
            num_id: Numbering definition ID
            level: List level
        """
        # Get or create paragraph properties
        p = paragraph._p
        pPr = p.get_or_add_pPr()

        # Create numPr element
        numPr = OxmlElement('w:numPr')

        # Add level reference
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(level))
        numPr.append(ilvl)

        # Add numbering ID reference
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), str(num_id))
        numPr.append(numId)

        # Insert numPr at the beginning of pPr
        pPr.insert(0, numPr)


def parse_inline_formatting(text: str) -> list[TextRun]:
    """Parse text with markdown-style inline formatting into TextRuns.

    Supports:
    - **bold**
    - *italic*
    - __underline__
    - ~~strikethrough~~

    Args:
        text: Text with inline formatting markers

    Returns:
        List of TextRun objects with appropriate formatting
    """
    runs = []

    # Pattern for various formatting markers
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|__(.+?)__|~~(.+?)~~|([^*_~]+))'

    for match in re.finditer(pattern, text):
        full_match = match.group(0)

        if match.group(2):  # Bold: **text**
            runs.append(TextRun(
                text=match.group(2),
                formatting=InlineFormatting(bold=True)
            ))
        elif match.group(3):  # Italic: *text*
            runs.append(TextRun(
                text=match.group(3),
                formatting=InlineFormatting(italic=True)
            ))
        elif match.group(4):  # Underline: __text__
            runs.append(TextRun(
                text=match.group(4),
                formatting=InlineFormatting(underline=True)
            ))
        elif match.group(5):  # Strikethrough: ~~text~~
            runs.append(TextRun(
                text=match.group(5),
                formatting=InlineFormatting(strike=True)
            ))
        elif match.group(6):  # Plain text
            if match.group(6).strip() or runs:  # Keep if not just leading whitespace
                runs.append(TextRun(text=match.group(6)))

    return runs if runs else [TextRun(text=text)]


def create_list_item_with_formatting(
    text: str,
    level: int = 0,
    parse_markdown: bool = True
) -> ListItem:
    """Convenience function to create a ListItem with optional markdown parsing.

    Args:
        text: Item text, optionally with markdown formatting
        level: Nesting level
        parse_markdown: Whether to parse markdown-style formatting

    Returns:
        ListItem with appropriate formatting
    """
    if parse_markdown:
        runs = parse_inline_formatting(text)
        return ListItem(content="", level=level, runs=runs)
    else:
        return ListItem(content=text, level=level)
