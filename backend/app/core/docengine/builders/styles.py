"""Style Engine for the Unified Document Engine.

This module provides document-wide styling capabilities based on UIF
(Unified Interchange Format) styling configuration. It handles font settings,
heading styles, page layout, margins, and header/footer configuration.
"""

import logging
from typing import Optional, Union
from dataclasses import dataclass, field

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


@dataclass
class DocumentStyling:
    """Document styling configuration from UIF."""

    default_font: str = "Arial"
    default_font_size: int = 11  # points
    heading_font: str = "Arial"
    heading_1_size: int = 16
    heading_2_size: int = 14
    heading_3_size: int = 12
    heading_4_size: int = 11
    line_spacing: float = 1.15
    heading_1_bold: bool = True
    heading_2_bold: bool = True
    heading_3_bold: bool = True
    heading_4_bold: bool = False
    heading_1_color: Optional[str] = None  # hex color e.g., "000000"
    heading_2_color: Optional[str] = None
    heading_3_color: Optional[str] = None
    heading_4_color: Optional[str] = None


@dataclass
class PageSetup:
    """Page layout configuration from UIF."""

    page_width: float = 8.5  # inches (Letter)
    page_height: float = 11  # inches (Letter)
    margin_top: float = 1.0
    margin_bottom: float = 1.0
    margin_left: float = 1.0
    margin_right: float = 1.0

    @classmethod
    def letter(cls) -> "PageSetup":
        """Return Letter size page setup (8.5 x 11 inches)."""
        return cls(page_width=8.5, page_height=11)

    @classmethod
    def a4(cls) -> "PageSetup":
        """Return A4 size page setup (8.27 x 11.69 inches)."""
        return cls(page_width=8.27, page_height=11.69)


@dataclass
class HeaderFooter:
    """Header and footer configuration from UIF."""

    header_text: Optional[str] = None
    footer_text: Optional[str] = None
    show_page_numbers: bool = True
    page_number_position: str = "footer_right"  # footer_center, footer_right, header_right
    header_font_size: int = 9
    footer_font_size: int = 9
    include_total_pages: bool = False  # Show "Page X of Y" format


class StyleEngine:
    """
    Engine for applying document-wide styling based on UIF configuration.

    This engine handles:
    - Default font family and size
    - Heading styles (Heading 1-4)
    - Body text style (Normal)
    - Page margins and size
    - Headers and footers with optional page numbers
    - Line spacing

    Example usage:
        doc = Document()
        engine = StyleEngine(doc)

        styling = {
            "default_font": "Calibri",
            "default_font_size": 11,
            "heading_font": "Calibri",
            "heading_1_size": 18,
            "line_spacing": 1.5
        }

        page_setup = {
            "page_width": 8.5,
            "page_height": 11,
            "margin_top": 1.0,
            "margin_bottom": 1.0
        }

        engine.apply_document_styling(styling)
        engine.setup_page_layout(page_setup)
        engine.add_header("Protocol XYZ-001")
        engine.add_footer("CONFIDENTIAL", show_page_number=True)
    """

    def __init__(self, doc: Document):
        """
        Initialize the StyleEngine with a python-docx Document.

        Args:
            doc: The python-docx Document object to apply styles to.
        """
        self.doc = doc
        self._ensure_styles_exist()

    def _ensure_styles_exist(self) -> None:
        """Ensure required styles exist in the document."""
        styles = self.doc.styles

        # Check and create heading styles if they don't exist
        for i in range(1, 5):
            style_name = f"Heading {i}"
            try:
                _ = styles[style_name]
            except KeyError:
                logger.debug(f"Creating style: {style_name}")
                styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        # Ensure Normal style exists (it should always exist)
        try:
            _ = styles["Normal"]
        except KeyError:
            logger.warning("Normal style not found, creating")
            styles.add_style("Normal", WD_STYLE_TYPE.PARAGRAPH)

    def apply_document_styling(self, styling: Union[dict, DocumentStyling]) -> None:
        """
        Apply full document styling from UIF configuration.

        This method configures all document-wide styles including default font,
        heading styles, and body text style.

        Args:
            styling: Dictionary or DocumentStyling object with styling configuration.
                Expected keys:
                - default_font: str (default: "Arial")
                - default_font_size: int in points (default: 11)
                - heading_font: str (default: "Arial")
                - heading_1_size through heading_4_size: int in points
                - line_spacing: float (default: 1.15)
                - heading_X_bold: bool (default: True for 1-3, False for 4)
                - heading_X_color: str hex color (optional)
        """
        if isinstance(styling, DocumentStyling):
            config = styling
        else:
            config = DocumentStyling(**styling) if styling else DocumentStyling()

        logger.info(f"Applying document styling: font={config.default_font}, size={config.default_font_size}pt")

        # Configure body text style first (other styles may inherit from it)
        self.configure_body_style({
            "default_font": config.default_font,
            "default_font_size": config.default_font_size,
            "line_spacing": config.line_spacing,
        })

        # Configure heading styles
        self.configure_heading_styles({
            "heading_font": config.heading_font,
            "heading_1_size": config.heading_1_size,
            "heading_2_size": config.heading_2_size,
            "heading_3_size": config.heading_3_size,
            "heading_4_size": config.heading_4_size,
            "heading_1_bold": config.heading_1_bold,
            "heading_2_bold": config.heading_2_bold,
            "heading_3_bold": config.heading_3_bold,
            "heading_4_bold": config.heading_4_bold,
            "heading_1_color": config.heading_1_color,
            "heading_2_color": config.heading_2_color,
            "heading_3_color": config.heading_3_color,
            "heading_4_color": config.heading_4_color,
            "line_spacing": config.line_spacing,
        })

    def setup_page_layout(self, page_setup: Union[dict, PageSetup]) -> None:
        """
        Configure page size and margins for all sections.

        Args:
            page_setup: Dictionary or PageSetup object with page configuration.
                Expected keys:
                - page_width: float in inches (default: 8.5 for Letter)
                - page_height: float in inches (default: 11 for Letter)
                - margin_top: float in inches (default: 1.0)
                - margin_bottom: float in inches (default: 1.0)
                - margin_left: float in inches (default: 1.0)
                - margin_right: float in inches (default: 1.0)
        """
        if isinstance(page_setup, PageSetup):
            config = page_setup
        else:
            config = PageSetup(**page_setup) if page_setup else PageSetup()

        logger.info(
            f"Setting up page layout: {config.page_width}x{config.page_height} inches, "
            f"margins: T={config.margin_top}, B={config.margin_bottom}, "
            f"L={config.margin_left}, R={config.margin_right}"
        )

        # Apply to all sections in the document
        for section in self.doc.sections:
            # Page size
            section.page_width = Inches(config.page_width)
            section.page_height = Inches(config.page_height)

            # Margins
            section.top_margin = Inches(config.margin_top)
            section.bottom_margin = Inches(config.margin_bottom)
            section.left_margin = Inches(config.margin_left)
            section.right_margin = Inches(config.margin_right)

    def configure_heading_styles(self, styling: dict) -> None:
        """
        Configure heading 1-4 styles with font, size, and formatting.

        Args:
            styling: Dictionary with heading style configuration.
                Expected keys:
                - heading_font: str (font family name)
                - heading_1_size through heading_4_size: int (font size in points)
                - heading_1_bold through heading_4_bold: bool (optional)
                - heading_1_color through heading_4_color: str hex color (optional)
                - line_spacing: float (optional)
        """
        heading_font = styling.get("heading_font", "Arial")
        line_spacing = styling.get("line_spacing", 1.15)

        heading_configs = [
            (1, styling.get("heading_1_size", 16), styling.get("heading_1_bold", True), styling.get("heading_1_color")),
            (2, styling.get("heading_2_size", 14), styling.get("heading_2_bold", True), styling.get("heading_2_color")),
            (3, styling.get("heading_3_size", 12), styling.get("heading_3_bold", True), styling.get("heading_3_color")),
            (4, styling.get("heading_4_size", 11), styling.get("heading_4_bold", False), styling.get("heading_4_color")),
        ]

        for level, size, bold, color in heading_configs:
            style_name = f"Heading {level}"
            try:
                style = self.doc.styles[style_name]
                self._set_style_font(style, heading_font, size)

                # Set bold
                style.font.bold = bold

                # Set color if specified
                if color:
                    self._set_font_color(style, color)

                # Set paragraph formatting
                self._set_paragraph_spacing(style, line_spacing, space_before=Pt(12), space_after=Pt(6))

                logger.debug(f"Configured {style_name}: {heading_font} {size}pt, bold={bold}")

            except KeyError:
                logger.warning(f"Style '{style_name}' not found in document")

    def configure_body_style(self, styling: dict) -> None:
        """
        Configure Normal/body text style.

        Args:
            styling: Dictionary with body style configuration.
                Expected keys:
                - default_font: str (font family name)
                - default_font_size: int (font size in points)
                - line_spacing: float (line spacing multiplier)
        """
        default_font = styling.get("default_font", "Arial")
        default_size = styling.get("default_font_size", 11)
        line_spacing = styling.get("line_spacing", 1.15)

        try:
            normal_style = self.doc.styles["Normal"]
            self._set_style_font(normal_style, default_font, default_size)
            self._set_paragraph_spacing(normal_style, line_spacing)

            logger.debug(f"Configured Normal style: {default_font} {default_size}pt, spacing={line_spacing}")

        except KeyError:
            logger.warning("Normal style not found in document")

    def add_header(self, text: str, alignment: str = "right") -> None:
        """
        Add document header to all sections.

        Args:
            text: Text to display in the header.
            alignment: Text alignment - "left", "center", or "right" (default: "right").
        """
        if not text:
            return

        logger.info(f"Adding header: '{text[:50]}...' aligned {alignment}")

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }

        for section in self.doc.sections:
            header = section.header
            header.is_linked_to_previous = False

            # Clear existing header content
            for paragraph in header.paragraphs:
                paragraph.clear()

            # Add header text
            if header.paragraphs:
                paragraph = header.paragraphs[0]
            else:
                paragraph = header.add_paragraph()

            run = paragraph.add_run(text)
            run.font.size = Pt(9)
            run.font.name = "Arial"
            paragraph.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.RIGHT)

    def add_footer(
        self,
        text: Optional[str] = None,
        show_page_number: bool = True,
        page_number_position: str = "footer_right",
        include_total_pages: bool = False
    ) -> None:
        """
        Add document footer with optional page numbers.

        Args:
            text: Text to display in the footer (optional).
            show_page_number: Whether to show page numbers (default: True).
            page_number_position: Position for page numbers - "footer_center",
                "footer_right", or "header_right" (default: "footer_right").
            include_total_pages: Show "Page X of Y" format (default: False).
        """
        logger.info(f"Adding footer: text='{text}', page_numbers={show_page_number}")

        for section in self.doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False

            # Clear existing footer content
            for paragraph in footer.paragraphs:
                paragraph.clear()

            # Determine alignment based on position
            if page_number_position == "footer_center":
                alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif page_number_position == "footer_right":
                alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Get or create paragraph
            if footer.paragraphs:
                paragraph = footer.paragraphs[0]
            else:
                paragraph = footer.add_paragraph()

            paragraph.alignment = alignment

            # Add footer text if provided
            if text:
                run = paragraph.add_run(text)
                run.font.size = Pt(9)
                run.font.name = "Arial"

                if show_page_number:
                    # Add separator before page number
                    sep_run = paragraph.add_run("  |  ")
                    sep_run.font.size = Pt(9)
                    sep_run.font.name = "Arial"

            # Add page number
            if show_page_number:
                self._add_page_number_field(paragraph, include_total_pages)

    def apply_header_footer_config(self, config: Union[dict, HeaderFooter]) -> None:
        """
        Apply header and footer configuration from UIF config.

        Args:
            config: Dictionary or HeaderFooter object with configuration.
                Expected keys:
                - header_text: str (optional)
                - footer_text: str (optional)
                - show_page_numbers: bool (default: True)
                - page_number_position: str (default: "footer_right")
                - include_total_pages: bool (default: False)
        """
        if isinstance(config, HeaderFooter):
            hf_config = config
        else:
            hf_config = HeaderFooter(**config) if config else HeaderFooter()

        if hf_config.header_text:
            # Determine header alignment from page number position
            if hf_config.page_number_position == "header_right":
                self.add_header(hf_config.header_text, alignment="left")
            else:
                self.add_header(hf_config.header_text, alignment="right")

        # Add footer with page numbers
        self.add_footer(
            text=hf_config.footer_text,
            show_page_number=hf_config.show_page_numbers,
            page_number_position=hf_config.page_number_position,
            include_total_pages=hf_config.include_total_pages
        )

        # If page numbers should be in header
        if hf_config.show_page_numbers and hf_config.page_number_position == "header_right":
            self._add_page_number_to_header(hf_config.include_total_pages)

    def _set_style_font(self, style, font_name: str, font_size: int) -> None:
        """
        Set font properties on a style.

        Args:
            style: The python-docx style object.
            font_name: Name of the font family (e.g., "Arial", "Calibri").
            font_size: Font size in points.
        """
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)

        # Set font for East Asian text (ensures consistent rendering)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def _set_font_color(self, style, hex_color: str) -> None:
        """
        Set font color on a style.

        Args:
            style: The python-docx style object.
            hex_color: Hex color string (e.g., "000000" for black, "0066CC" for blue).
        """
        try:
            # Parse hex color
            hex_color = hex_color.lstrip('#')
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)

            style.font.color.rgb = RGBColor(r, g, b)
        except (ValueError, IndexError) as e:
            logger.warning(f"Invalid hex color '{hex_color}': {e}")

    def _set_paragraph_spacing(
        self,
        style,
        line_spacing: float,
        space_before: Optional[Pt] = None,
        space_after: Optional[Pt] = None
    ) -> None:
        """
        Set paragraph spacing properties on a style.

        Args:
            style: The python-docx style object.
            line_spacing: Line spacing multiplier (e.g., 1.0, 1.15, 1.5, 2.0).
            space_before: Space before paragraph in points (optional).
            space_after: Space after paragraph in points (optional).
        """
        paragraph_format = style.paragraph_format

        # Set line spacing
        paragraph_format.line_spacing = line_spacing
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

        # Set space before/after if specified
        if space_before is not None:
            paragraph_format.space_before = space_before
        if space_after is not None:
            paragraph_format.space_after = space_after

    def _add_page_number_field(self, paragraph, include_total: bool = False) -> None:
        """
        Add a page number field to a paragraph.

        Args:
            paragraph: The python-docx paragraph object.
            include_total: Whether to include total pages ("Page X of Y").
        """
        run = paragraph.add_run()
        run.font.size = Pt(9)
        run.font.name = "Arial"

        # Add "Page " prefix
        fld_prefix = OxmlElement('w:r')
        fld_prefix_text = OxmlElement('w:t')
        fld_prefix_text.text = "Page "
        fld_prefix.append(fld_prefix_text)
        run._r.append(fld_prefix)

        # Add PAGE field
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = " PAGE "

        fld_char_separate = OxmlElement('w:fldChar')
        fld_char_separate.set(qn('w:fldCharType'), 'separate')

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)
        run._r.append(fld_char_end)

        if include_total:
            # Add " of " text
            run2 = paragraph.add_run(" of ")
            run2.font.size = Pt(9)
            run2.font.name = "Arial"

            # Add NUMPAGES field for total
            run3 = paragraph.add_run()
            run3.font.size = Pt(9)
            run3.font.name = "Arial"

            fld_char_begin2 = OxmlElement('w:fldChar')
            fld_char_begin2.set(qn('w:fldCharType'), 'begin')

            instr_text2 = OxmlElement('w:instrText')
            instr_text2.set(qn('xml:space'), 'preserve')
            instr_text2.text = " NUMPAGES "

            fld_char_separate2 = OxmlElement('w:fldChar')
            fld_char_separate2.set(qn('w:fldCharType'), 'separate')

            fld_char_end2 = OxmlElement('w:fldChar')
            fld_char_end2.set(qn('w:fldCharType'), 'end')

            run3._r.append(fld_char_begin2)
            run3._r.append(instr_text2)
            run3._r.append(fld_char_separate2)
            run3._r.append(fld_char_end2)

    def _add_page_number_to_header(self, include_total: bool = False) -> None:
        """
        Add page number to the right side of the header.

        Args:
            include_total: Whether to include total pages.
        """
        for section in self.doc.sections:
            header = section.header

            # If header already has content, we need to handle alignment
            # Create a table for left/right alignment
            if header.paragraphs and header.paragraphs[0].text:
                # Header has existing text, add page number to right
                existing_text = header.paragraphs[0].text
                header.paragraphs[0].clear()

                # Create a table with two cells for left/right alignment
                table = header.add_table(rows=1, cols=2)
                table.autofit = True
                table.allow_autofit = True

                # Left cell - existing text
                left_cell = table.rows[0].cells[0]
                left_para = left_cell.paragraphs[0]
                left_run = left_para.add_run(existing_text)
                left_run.font.size = Pt(9)
                left_run.font.name = "Arial"

                # Right cell - page number
                right_cell = table.rows[0].cells[1]
                right_para = right_cell.paragraphs[0]
                right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                self._add_page_number_field(right_para, include_total)
            else:
                # No existing text, just add page number
                if header.paragraphs:
                    paragraph = header.paragraphs[0]
                else:
                    paragraph = header.add_paragraph()

                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                self._add_page_number_field(paragraph, include_total)

    def set_default_paragraph_style(self, style_name: str = "Normal") -> None:
        """
        Set the default paragraph style for new paragraphs.

        Args:
            style_name: Name of the style to use as default.
        """
        try:
            self.doc.styles.default(WD_STYLE_TYPE.PARAGRAPH).name = style_name
            logger.debug(f"Set default paragraph style to: {style_name}")
        except Exception as e:
            logger.warning(f"Could not set default paragraph style: {e}")

    def create_custom_style(
        self,
        name: str,
        font_name: str,
        font_size: int,
        bold: bool = False,
        italic: bool = False,
        color: Optional[str] = None,
        line_spacing: float = 1.15
    ) -> None:
        """
        Create a custom paragraph style.

        Args:
            name: Name for the new style.
            font_name: Font family name.
            font_size: Font size in points.
            bold: Whether text should be bold.
            italic: Whether text should be italic.
            color: Hex color string (optional).
            line_spacing: Line spacing multiplier.
        """
        try:
            # Check if style already exists
            try:
                style = self.doc.styles[name]
                logger.debug(f"Style '{name}' already exists, updating")
            except KeyError:
                style = self.doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
                logger.debug(f"Created new style: {name}")

            self._set_style_font(style, font_name, font_size)
            style.font.bold = bold
            style.font.italic = italic

            if color:
                self._set_font_color(style, color)

            self._set_paragraph_spacing(style, line_spacing)

        except Exception as e:
            logger.error(f"Failed to create custom style '{name}': {e}")
            raise


def apply_full_styling(
    doc: Document,
    styling: Union[dict, DocumentStyling, None] = None,
    page_setup: Union[dict, PageSetup, None] = None,
    header_footer: Union[dict, HeaderFooter, None] = None
) -> StyleEngine:
    """
    Convenience function to apply full document styling.

    Creates a StyleEngine and applies all styling configurations in one call.

    Args:
        doc: The python-docx Document object.
        styling: Document styling configuration (fonts, sizes).
        page_setup: Page layout configuration (size, margins).
        header_footer: Header and footer configuration.

    Returns:
        The configured StyleEngine instance for further customization.

    Example:
        from docx import Document
        from app.core.docengine.builders.styles import apply_full_styling

        doc = Document()
        engine = apply_full_styling(
            doc,
            styling={"default_font": "Calibri", "heading_1_size": 18},
            page_setup={"margin_top": 1.5, "margin_bottom": 1.5},
            header_footer={"header_text": "CONFIDENTIAL", "show_page_numbers": True}
        )
    """
    engine = StyleEngine(doc)

    if styling:
        engine.apply_document_styling(styling)
    else:
        # Apply default styling
        engine.apply_document_styling(DocumentStyling())

    if page_setup:
        engine.setup_page_layout(page_setup)
    else:
        # Apply default page setup
        engine.setup_page_layout(PageSetup())

    if header_footer:
        engine.apply_header_footer_config(header_footer)

    return engine
