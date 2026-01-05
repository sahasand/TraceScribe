"""Section and Paragraph builder for the Unified Document Engine."""

import logging
from typing import List, Optional, Dict, Union
from dataclasses import dataclass, field
from enum import Enum

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.text.run import Run

logger = logging.getLogger(__name__)


class Alignment(str, Enum):
    """Text alignment options."""
    LEFT = "left"
    CENTER = "center"
    RIGHT = "right"
    JUSTIFY = "justify"


class ContentBlockType(str, Enum):
    """Types of content blocks."""
    PARAGRAPH = "paragraph"
    HEADING = "heading"
    PAGE_BREAK = "page_break"


@dataclass
class InlineFormat:
    """Defines formatting for a range of characters."""
    start: int
    end: int
    bold: bool = False
    italic: bool = False
    underline: bool = False


@dataclass
class InlineFormatting:
    """Container for inline formatting ranges."""
    ranges: List[InlineFormat] = field(default_factory=list)


@dataclass
class ContentBlock:
    """A block of content within a section."""
    type: str
    content: str = ""
    level: int = 1
    formatting: Optional[InlineFormatting] = None
    alignment: str = "left"
    spacing_before: int = 0
    spacing_after: int = 0


@dataclass
class Section:
    """A document section with heading and content."""
    id: str
    level: int
    heading: str
    content_blocks: List[ContentBlock] = field(default_factory=list)
    subsections: List["Section"] = field(default_factory=list)


class SectionBuilder:
    """
    Builder for rendering UIF sections and paragraphs to python-docx.

    Handles:
    - Heading levels 1-4
    - Paragraphs with inline formatting (bold, italic, underline)
    - Page breaks
    - Text alignment (left, center, right, justify)
    - Paragraph spacing (before/after in points)
    """

    ALIGNMENT_MAP = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    DEFAULT_HEADING_SPACING_BEFORE = {1: 24, 2: 18, 3: 12, 4: 12}
    DEFAULT_HEADING_SPACING_AFTER = {1: 12, 2: 6, 3: 6, 4: 6}

    def __init__(self, doc: Document):
        """
        Initialize SectionBuilder.

        Args:
            doc: python-docx Document instance to build into
        """
        self.doc = doc

    def add_section(self, section: Union[dict, Section]) -> None:
        """
        Add a complete section to the document.

        Renders the section heading, all content blocks, and recursively
        processes any subsections.

        Args:
            section: Section data as dict or Section dataclass
        """
        if isinstance(section, Section):
            section = self._section_to_dict(section)

        if not isinstance(section, dict):
            logger.warning(f"Invalid section type: {type(section)}")
            return

        heading = section.get("heading", "").strip()
        level = section.get("level", 1)

        if heading:
            self.add_heading(heading, level)

        for block in section.get("content_blocks", []):
            self._process_content_block(block)

        for subsection in section.get("subsections", []):
            self.add_section(subsection)

    def add_heading(self, text: str, level: int) -> Paragraph:
        """
        Add a heading to the document.

        Args:
            text: Heading text
            level: Heading level (1-4, clamped to valid range)

        Returns:
            The created Paragraph object
        """
        text = str(text).strip() if text else ""
        if not text:
            return self.doc.add_paragraph("")

        level = max(1, min(4, int(level)))

        paragraph = self.doc.add_heading(text, level=level)

        pf = paragraph.paragraph_format
        pf.space_before = Pt(self.DEFAULT_HEADING_SPACING_BEFORE.get(level, 12))
        pf.space_after = Pt(self.DEFAULT_HEADING_SPACING_AFTER.get(level, 6))

        return paragraph

    def add_paragraph(self, block: Union[dict, ContentBlock]) -> Paragraph:
        """
        Add a paragraph to the document.

        Args:
            block: Content block data as dict or ContentBlock dataclass

        Returns:
            The created Paragraph object
        """
        if isinstance(block, ContentBlock):
            block = self._content_block_to_dict(block)

        content = str(block.get("content", "")).strip() if block.get("content") else ""

        paragraph = self.doc.add_paragraph()

        alignment = block.get("alignment", "left").lower()
        paragraph.alignment = self.ALIGNMENT_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

        pf = paragraph.paragraph_format
        spacing_before = block.get("spacing_before", 0)
        spacing_after = block.get("spacing_after", 0)

        if spacing_before and spacing_before > 0:
            pf.space_before = Pt(spacing_before)
        if spacing_after and spacing_after > 0:
            pf.space_after = Pt(spacing_after)

        formatting = block.get("formatting")
        if formatting and content:
            self._apply_inline_formatting(paragraph, content, formatting)
        elif content:
            paragraph.add_run(content)

        return paragraph

    def add_page_break(self) -> None:
        """Add a page break to the document."""
        self.doc.add_page_break()

    def _process_content_block(self, block: Union[dict, ContentBlock]) -> Optional[Paragraph]:
        """
        Process a content block based on its type.

        Args:
            block: Content block data

        Returns:
            The created Paragraph object, or None for page breaks
        """
        if isinstance(block, ContentBlock):
            block = self._content_block_to_dict(block)

        if not isinstance(block, dict):
            logger.warning(f"Invalid content block type: {type(block)}")
            return None

        block_type = str(block.get("type", "paragraph")).lower()

        if block_type == "paragraph":
            return self.add_paragraph(block)
        elif block_type == "heading":
            return self.add_heading(block.get("content", ""), block.get("level", 1))
        elif block_type == "page_break":
            self.add_page_break()
            return None
        else:
            logger.warning(f"Unknown content block type: {block_type}, treating as paragraph")
            return self.add_paragraph(block)

    def _apply_inline_formatting(
        self,
        paragraph: Paragraph,
        content: str,
        formatting: Union[dict, InlineFormatting]
    ) -> None:
        """
        Apply inline formatting to paragraph content.

        Handles overlapping formatting ranges by building a character-level
        format map and creating runs for each unique format combination.

        Args:
            paragraph: Target paragraph
            content: Text content
            formatting: Formatting specification with ranges
        """
        if isinstance(formatting, InlineFormatting):
            formatting = {
                "ranges": [self._inline_format_to_dict(r) for r in formatting.ranges]
            }

        ranges = formatting.get("ranges", [])
        if not ranges:
            paragraph.add_run(content)
            return

        validated_ranges = self._validate_and_sort_ranges(ranges, len(content))
        if not validated_ranges:
            paragraph.add_run(content)
            return

        char_formats = self._build_character_format_map(validated_ranges, len(content))
        self._create_formatted_runs(paragraph, content, char_formats)

    def _validate_and_sort_ranges(self, ranges: List[dict], content_length: int) -> List[dict]:
        """
        Validate and sort formatting ranges.

        Args:
            ranges: List of formatting range dicts
            content_length: Length of content string

        Returns:
            Validated and sorted ranges
        """
        validated = []

        for range_dict in ranges:
            if not isinstance(range_dict, dict):
                continue

            try:
                start = max(0, min(int(range_dict.get("start", 0)), content_length))
                end = max(start, min(int(range_dict.get("end", 0)), content_length))

                if start >= end:
                    continue

                validated.append({
                    "start": start,
                    "end": end,
                    "bold": bool(range_dict.get("bold", False)),
                    "italic": bool(range_dict.get("italic", False)),
                    "underline": bool(range_dict.get("underline", False)),
                })
            except (TypeError, ValueError) as e:
                logger.warning(f"Invalid formatting range: {range_dict}, error: {e}")
                continue

        validated.sort(key=lambda r: r["start"])
        return validated

    def _build_character_format_map(
        self,
        ranges: List[dict],
        content_length: int
    ) -> List[Dict[str, bool]]:
        """
        Build a per-character format map from formatting ranges.

        This handles overlapping ranges by applying formats cumulatively.

        Args:
            ranges: Validated formatting ranges
            content_length: Length of content string

        Returns:
            List of format dicts, one per character
        """
        char_formats = [
            {"bold": False, "italic": False, "underline": False}
            for _ in range(content_length)
        ]

        for range_dict in ranges:
            for i in range(range_dict["start"], range_dict["end"]):
                if range_dict.get("bold"):
                    char_formats[i]["bold"] = True
                if range_dict.get("italic"):
                    char_formats[i]["italic"] = True
                if range_dict.get("underline"):
                    char_formats[i]["underline"] = True

        return char_formats

    def _create_formatted_runs(
        self,
        paragraph: Paragraph,
        content: str,
        char_formats: List[Dict[str, bool]]
    ) -> None:
        """
        Create runs with appropriate formatting from character format map.

        Consolidates consecutive characters with identical formatting into
        single runs for efficiency.

        Args:
            paragraph: Target paragraph
            content: Text content
            char_formats: Per-character format map
        """
        if not content or not char_formats:
            return

        current_start = 0
        current_format = char_formats[0]

        for i in range(1, len(content)):
            if char_formats[i] != current_format:
                self._add_formatted_run(paragraph, content[current_start:i], current_format)
                current_start = i
                current_format = char_formats[i]

        self._add_formatted_run(paragraph, content[current_start:], current_format)

    def _add_formatted_run(
        self,
        paragraph: Paragraph,
        text: str,
        format_spec: Dict[str, bool]
    ) -> Run:
        """
        Add a formatted run to a paragraph.

        Args:
            paragraph: Target paragraph
            text: Run text
            format_spec: Format specification dict

        Returns:
            The created Run object
        """
        run = paragraph.add_run(text)
        run.bold = format_spec.get("bold", False)
        run.italic = format_spec.get("italic", False)
        run.underline = format_spec.get("underline", False)
        return run

    def _section_to_dict(self, section: Section) -> dict:
        """Convert Section dataclass to dict."""
        return {
            "id": section.id,
            "level": section.level,
            "heading": section.heading,
            "content_blocks": [
                self._content_block_to_dict(b) for b in section.content_blocks
            ],
            "subsections": [
                self._section_to_dict(s) for s in section.subsections
            ],
        }

    def _content_block_to_dict(self, block: ContentBlock) -> dict:
        """Convert ContentBlock dataclass to dict."""
        result = {
            "type": block.type,
            "content": block.content,
            "level": block.level,
            "alignment": block.alignment,
            "spacing_before": block.spacing_before,
            "spacing_after": block.spacing_after,
        }
        if block.formatting:
            result["formatting"] = {
                "ranges": [
                    self._inline_format_to_dict(f) for f in block.formatting.ranges
                ]
            }
        return result

    def _inline_format_to_dict(self, fmt: InlineFormat) -> dict:
        """Convert InlineFormat dataclass to dict."""
        return {
            "start": fmt.start,
            "end": fmt.end,
            "bold": fmt.bold,
            "italic": fmt.italic,
            "underline": fmt.underline,
        }

    def add_sections(self, sections: List[Union[dict, Section]]) -> None:
        """
        Add multiple sections to the document.

        Args:
            sections: List of section data (dicts or Section dataclasses)
        """
        for section in sections:
            self.add_section(section)
