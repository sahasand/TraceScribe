"""
Table Builder for the Unified Document Engine.

This module handles rendering UIF table blocks to python-docx tables,
supporting headers, data rows, column widths, cell styling, and merged cells.
"""

import logging
from typing import Any, Dict, List, Optional, Union

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell

logger = logging.getLogger(__name__)


class TableCell:
    """
    Represents a cell in a UIF table with styling options.

    Attributes:
        content: Text content of the cell.
        colspan: Number of columns this cell spans (default 1).
        rowspan: Number of rows this cell spans (default 1, not fully supported).
        alignment: Horizontal alignment (left, center, right).
        vertical_alignment: Vertical alignment (top, center, bottom).
        background_color: Hex color string (e.g., #CCCCCC) for cell shading.
    """

    def __init__(
        self,
        content: str = "",
        colspan: int = 1,
        rowspan: int = 1,
        alignment: str = "left",
        vertical_alignment: str = "center",
        background_color: Optional[str] = None,
    ):
        self.content = content
        self.colspan = max(1, colspan)
        self.rowspan = max(1, rowspan)
        self.alignment = alignment.lower()
        self.vertical_alignment = vertical_alignment.lower()
        self.background_color = background_color


class TableBuilder:
    """
    Builder for rendering UIF table blocks to python-docx Document tables.

    This builder supports:
    - Header rows with distinct background styling
    - Data rows with string or TableCell content
    - Column widths (auto or specified in inches)
    - Cell borders (via table style)
    - Cell background colors/shading
    - Cell alignment (horizontal and vertical)
    - Merged cells (colspan)

    Example usage:
        from docx import Document
        from app.core.docengine.builders.table import TableBuilder

        doc = Document()
        builder = TableBuilder(doc)

        table_block = {
            "type": "table",
            "headers": ["Name", "Age", "City"],
            "rows": [
                ["Alice", "30", "New York"],
                ["Bob", "25", "Los Angeles"],
            ],
            "column_widths": [2.0, 1.0, 2.0],
            "style": "Table Grid",
            "header_background": "#4472C4",
        }

        builder.add_table(table_block)
        doc.save("output.docx")
    """

    # Mapping from string alignment to python-docx enums
    HORIZONTAL_ALIGNMENT_MAP = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    VERTICAL_ALIGNMENT_MAP = {
        "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
        "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
    }

    def __init__(self, doc: Document):
        """
        Initialize the TableBuilder with a python-docx Document.

        Args:
            doc: The python-docx Document instance to add tables to.
        """
        self.doc = doc

    def add_table(self, block: Dict[str, Any]) -> Optional[Table]:
        """
        Add a table to the document from UIF block data.

        Args:
            block: A dictionary containing table configuration:
                - type (str): Should be "table"
                - headers (List[str]): Column header texts
                - rows (List[List[str | dict | TableCell]]): Data rows
                - column_widths (Optional[List[float]]): Column widths in inches
                - style (str): Table style name (default: "Table Grid")
                - header_background (str): Hex color for header background

        Returns:
            The created Table object, or None if the table could not be created.

        Raises:
            ValueError: If block is missing required fields.
        """
        if not block:
            logger.warning("Empty table block provided")
            return None

        headers = block.get("headers", [])
        rows = block.get("rows", [])

        # Handle empty table case
        if not headers and not rows:
            logger.warning("Table block has no headers or rows")
            return None

        # Determine number of columns
        num_cols = len(headers) if headers else self._get_max_columns(rows)
        if num_cols == 0:
            logger.warning("Could not determine number of columns for table")
            return None

        # Determine number of rows (headers + data rows)
        num_rows = (1 if headers else 0) + len(rows)
        if num_rows == 0:
            logger.warning("Table has no content rows")
            return None

        # Create the table
        table = self.doc.add_table(rows=num_rows, cols=num_cols)

        # Apply table style
        style = block.get("style", "Table Grid")
        try:
            table.style = style
        except KeyError:
            logger.warning(f"Table style '{style}' not found, using default")
            try:
                table.style = "Table Grid"
            except KeyError:
                pass  # Use whatever default is available

        # Set table alignment to center in document
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Track current row index
        row_idx = 0

        # Add header row if present
        if headers:
            header_background = block.get("header_background", "#CCCCCC")
            self._add_header_row(table, headers, header_background, row_idx)
            row_idx += 1

        # Add data rows
        for row_data in rows:
            self._add_data_row(table, row_data, row_idx, num_cols)
            row_idx += 1

        # Set column widths if specified
        column_widths = block.get("column_widths")
        if column_widths:
            self._set_column_widths(table, column_widths)

        # Add spacing after table
        self.doc.add_paragraph()

        return table

    def _get_max_columns(self, rows: List[List[Any]]) -> int:
        """
        Determine the maximum number of columns from row data.

        Takes into account colspan when calculating column count.

        Args:
            rows: List of row data.

        Returns:
            Maximum number of columns needed.
        """
        max_cols = 0
        for row in rows:
            if not row:
                continue
            col_count = 0
            for cell in row:
                if isinstance(cell, TableCell):
                    col_count += cell.colspan
                elif isinstance(cell, dict):
                    col_count += cell.get("colspan", 1)
                else:
                    col_count += 1
            max_cols = max(max_cols, col_count)
        return max_cols

    def _add_header_row(
        self,
        table: Table,
        headers: List[str],
        background: str,
        row_idx: int = 0,
    ) -> None:
        """
        Add a styled header row to the table.

        Args:
            table: The python-docx Table object.
            headers: List of header text strings.
            background: Hex color string for header background.
            row_idx: Row index in the table (typically 0).
        """
        if row_idx >= len(table.rows):
            logger.error(f"Row index {row_idx} out of range")
            return

        row = table.rows[row_idx]

        for col_idx, header_text in enumerate(headers):
            if col_idx >= len(row.cells):
                logger.warning(f"Header column {col_idx} exceeds table columns")
                break

            cell = row.cells[col_idx]

            # Set cell text
            cell.text = str(header_text) if header_text is not None else ""

            # Style the paragraph
            if cell.paragraphs:
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Bold header text
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)

            # Set background color
            self._set_cell_shading(cell, background)

            # Center vertically
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _add_data_row(
        self,
        table: Table,
        row_data: List[Union[str, Dict[str, Any], TableCell]],
        row_idx: int,
        num_cols: int,
    ) -> None:
        """
        Add a data row to the table.

        Args:
            table: The python-docx Table object.
            row_data: List of cell content (strings, dicts, or TableCell objects).
            row_idx: Row index in the table.
            num_cols: Total number of columns in the table.
        """
        if row_idx >= len(table.rows):
            logger.error(f"Row index {row_idx} out of range")
            return

        row = table.rows[row_idx]
        col_idx = 0

        for cell_data in row_data:
            if col_idx >= num_cols:
                logger.warning("Row data exceeds number of columns")
                break

            if col_idx >= len(row.cells):
                break

            cell = row.cells[col_idx]

            # Parse cell data
            if isinstance(cell_data, TableCell):
                cell_config = cell_data
            elif isinstance(cell_data, dict):
                cell_config = TableCell(
                    content=cell_data.get("content", ""),
                    colspan=cell_data.get("colspan", 1),
                    rowspan=cell_data.get("rowspan", 1),
                    alignment=cell_data.get("alignment", "left"),
                    vertical_alignment=cell_data.get("vertical_alignment", "center"),
                    background_color=cell_data.get("background_color"),
                )
            else:
                # Plain string content
                cell_config = TableCell(content=str(cell_data) if cell_data is not None else "")

            # Handle colspan (merge cells horizontally)
            if cell_config.colspan > 1:
                merge_end_idx = min(col_idx + cell_config.colspan - 1, num_cols - 1)
                if merge_end_idx > col_idx:
                    try:
                        cell = cell.merge(row.cells[merge_end_idx])
                    except Exception as e:
                        logger.warning(f"Failed to merge cells: {e}")

            # Set cell content
            cell.text = cell_config.content

            # Apply horizontal alignment
            if cell.paragraphs:
                para = cell.paragraphs[0]
                h_align = self.HORIZONTAL_ALIGNMENT_MAP.get(
                    cell_config.alignment, WD_ALIGN_PARAGRAPH.LEFT
                )
                para.alignment = h_align

            # Apply vertical alignment
            v_align = self.VERTICAL_ALIGNMENT_MAP.get(
                cell_config.vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )
            cell.vertical_alignment = v_align

            # Apply background color if specified
            if cell_config.background_color:
                self._set_cell_shading(cell, cell_config.background_color)

            # Move column index forward by colspan
            col_idx += cell_config.colspan

        # Fill remaining cells with empty content if row is short
        while col_idx < num_cols and col_idx < len(row.cells):
            row.cells[col_idx].text = ""
            col_idx += 1

    def _set_cell_shading(self, cell: _Cell, color: str) -> None:
        """
        Set the background color (shading) of a table cell.

        Args:
            cell: The python-docx cell object.
            color: Hex color string (e.g., #CCCCCC or CCCCCC).
        """
        if not color:
            return

        # Remove # prefix if present
        color = color.lstrip("#")

        # Validate hex color format
        if len(color) not in (3, 6):
            logger.warning(f"Invalid color format: {color}")
            return

        # Expand 3-character hex to 6-character
        if len(color) == 3:
            color = "".join(c * 2 for c in color)

        # Validate hex characters
        try:
            int(color, 16)
        except ValueError:
            logger.warning(f"Invalid hex color: {color}")
            return

        try:
            # Create shading element
            shading_xml = (
                f'<w:shd {nsdecls("w")} '
                f'w:val="clear" '
                f'w:color="auto" '
                f'w:fill="{color.upper()}"/>'
            )
            shading_element = parse_xml(shading_xml)

            # Get or create tcPr (table cell properties)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # Remove existing shading if present
            existing_shd = tcPr.find(qn("w:shd"))
            if existing_shd is not None:
                tcPr.remove(existing_shd)

            # Add new shading
            tcPr.append(shading_element)

        except Exception as e:
            logger.warning(f"Failed to set cell shading: {e}")

    def _set_column_widths(self, table: Table, widths: List[float]) -> None:
        """
        Set the width of each column in the table.

        Args:
            table: The python-docx Table object.
            widths: List of column widths in inches.
        """
        if not widths or not table.rows:
            return

        # Ensure table uses fixed column widths
        table.autofit = False

        for row in table.rows:
            for col_idx, width in enumerate(widths):
                if col_idx >= len(row.cells):
                    break

                try:
                    row.cells[col_idx].width = Inches(width)
                except Exception as e:
                    logger.warning(f"Failed to set column {col_idx} width: {e}")

    def add_table_from_data(
        self,
        headers: List[str],
        rows: List[List[Any]],
        column_widths: Optional[List[float]] = None,
        style: str = "Table Grid",
        header_background: str = "#CCCCCC",
    ) -> Optional[Table]:
        """
        Convenience method to add a table without creating a block dict.

        Args:
            headers: Column header texts.
            rows: List of row data (strings or TableCell objects).
            column_widths: Optional list of column widths in inches.
            style: Table style name.
            header_background: Hex color for header background.

        Returns:
            The created Table object, or None if creation failed.
        """
        block = {
            "type": "table",
            "headers": headers,
            "rows": rows,
            "style": style,
            "header_background": header_background,
        }

        if column_widths:
            block["column_widths"] = column_widths

        return self.add_table(block)
