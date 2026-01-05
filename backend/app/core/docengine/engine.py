"""
Unified Document Engine for TraceScribe.

This module provides the main DocEngine class that orchestrates the rendering
of UIF (Universal Intermediate Format) documents to Word format.
"""

from __future__ import annotations

import logging
import os
import tempfile
import time
import uuid
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, List, Dict, Any

from docx import Document

from app.core.storage import storage
from .schema import (
    UniversalDocument,
    ContentBlock,
    ContentBlockType,
    Section,
    SignatureBlock,
    TableBlock,
    ListStyle,
)
from .builders.section import SectionBuilder
from .builders.table import TableBuilder
from .builders.list import ListBuilder
from .builders.styles import StyleEngine

logger = logging.getLogger(__name__)


@dataclass
class DocEngineResult:
    """Result of document generation."""

    file_path: str
    file_size: int
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ValidationError:
    """A single validation error."""

    path: str
    message: str
    severity: str = "error"  # "error" or "warning"


class DocEngine:
    """
    Unified Document Engine.

    Takes a UniversalDocument (UIF JSON schema) and renders
    it to a Word document using python-docx.

    This is the central orchestrator that coordinates:
    - SectionBuilder: Renders sections, paragraphs, and headings
    - TableBuilder: Renders tables with headers, rows, and styling
    - ListBuilder: Renders bullet and numbered lists
    - StyleEngine: Applies document-wide styling and page setup

    Example usage:
        from app.core.docengine.engine import doc_engine
        from app.core.docengine.schema import UniversalDocument

        # Create a UIF document
        uif = UniversalDocument(
            document_type="icf",
            title="Informed Consent Form",
            sections=[...]
        )

        # Render to Word
        path = doc_engine.render(uif)

        # Or render and upload to storage
        result = await doc_engine.render_and_upload(
            uif,
            user_id="user123",
            protocol_id="proto456",
            version=1
        )
    """

    def __init__(self):
        """Initialize the document engine."""
        self.doc: Optional[Document] = None
        self.section_builder: Optional[SectionBuilder] = None
        self.table_builder: Optional[TableBuilder] = None
        self.list_builder: Optional[ListBuilder] = None
        self.style_engine: Optional[StyleEngine] = None

    def _init_document(self) -> None:
        """Initialize a new Word document with all builders."""
        self.doc = Document()
        self.section_builder = SectionBuilder(self.doc)
        self.table_builder = TableBuilder(self.doc)
        self.list_builder = ListBuilder(self.doc)
        self.style_engine = StyleEngine(self.doc)

    def validate(self, document: UniversalDocument) -> tuple[bool, List[str]]:
        """
        Validate a UniversalDocument for required fields and structure.

        Args:
            document: The UIF document to validate

        Returns:
            Tuple of (is_valid, list of error messages)
        """
        errors: List[str] = []

        # Required top-level fields
        if not document.title:
            errors.append("Document title is required")

        if not document.document_type:
            errors.append("Document type is required")

        if not document.sections:
            errors.append("Document must have at least one section")

        # Validate document type
        valid_types = {"icf", "dmp", "sap"}
        if document.document_type and document.document_type.lower() not in valid_types:
            errors.append(
                f"Invalid document type '{document.document_type}'. "
                f"Must be one of: {', '.join(valid_types)}"
            )

        # Validate sections recursively
        for i, section in enumerate(document.sections):
            section_path = f"sections[{i}]"
            self._validate_section(section, section_path, errors)

        # Validate page setup
        if document.page_setup:
            ps = document.page_setup
            if ps.page_width <= 0:
                errors.append("Page width must be positive")
            if ps.page_height <= 0:
                errors.append("Page height must be positive")
            if ps.margin_top < 0:
                errors.append("Top margin cannot be negative")
            if ps.margin_bottom < 0:
                errors.append("Bottom margin cannot be negative")
            if ps.margin_left < 0:
                errors.append("Left margin cannot be negative")
            if ps.margin_right < 0:
                errors.append("Right margin cannot be negative")

        # Validate styling
        if document.styling:
            if document.styling.default_font_size <= 0:
                errors.append("Default font size must be positive")
            if document.styling.line_spacing <= 0:
                errors.append("Line spacing must be positive")

        return len(errors) == 0, errors

    def _validate_section(
        self, section: Section, path: str, errors: List[str]
    ) -> None:
        """
        Validate a section recursively.

        Args:
            section: Section to validate
            path: Path string for error messages
            errors: List to append errors to
        """
        if not section.id:
            errors.append(f"{path}: Section ID is required")

        if not section.heading:
            errors.append(f"{path}: Section heading is required")

        if section.level < 1 or section.level > 4:
            errors.append(
                f"{path}: Section level must be between 1 and 4, got {section.level}"
            )

        # Validate content blocks
        for i, block in enumerate(section.content_blocks):
            block_path = f"{path}.content_blocks[{i}]"
            self._validate_content_block(block, block_path, errors)

        # Validate subsections recursively
        for i, subsection in enumerate(section.subsections):
            subsection_path = f"{path}.subsections[{i}]"

            # Check subsection level is greater than parent
            if subsection.level <= section.level:
                errors.append(
                    f"{subsection_path}: Subsection level ({subsection.level}) "
                    f"must be greater than parent level ({section.level})"
                )

            self._validate_section(subsection, subsection_path, errors)

    def _validate_content_block(
        self, block: ContentBlock, path: str, errors: List[str]
    ) -> None:
        """
        Validate a content block.

        Args:
            block: Content block to validate
            path: Path string for error messages
            errors: List to append errors to
        """
        # Type-specific validation
        if block.type == ContentBlockType.HEADING:
            if block.level is None or block.level < 1 or block.level > 4:
                errors.append(f"{path}: Heading level must be between 1 and 4")

        elif block.type == ContentBlockType.TABLE:
            if block.table is None:
                errors.append(f"{path}: Table block requires table data")
            elif not block.table.headers and not block.table.rows:
                errors.append(f"{path}: Table must have headers or rows")

        elif block.type in (ContentBlockType.BULLET_LIST, ContentBlockType.NUMBERED_LIST):
            if not block.items:
                errors.append(f"{path}: List block requires items")

        elif block.type == ContentBlockType.SIGNATURE_BLOCK:
            if block.signature is None:
                errors.append(f"{path}: Signature block requires signature data")

    def render(
        self,
        document: UniversalDocument,
        output_path: Optional[Path] = None,
        validate_first: bool = True,
    ) -> Path:
        """
        Render a UniversalDocument to a Word file.

        Args:
            document: The UIF document to render
            output_path: Optional output path (generates temp file if not provided)
            validate_first: Whether to validate the document before rendering

        Returns:
            Path to the generated Word file

        Raises:
            ValueError: If validation fails and validate_first is True
        """
        logger.info(f"Rendering document: {document.title}")

        # Optionally validate first
        if validate_first:
            is_valid, errors = self.validate(document)
            if not is_valid:
                error_msg = "Document validation failed:\n" + "\n".join(
                    f"  - {e}" for e in errors
                )
                logger.error(error_msg)
                raise ValueError(error_msg)

        # Initialize fresh document
        self._init_document()

        # Apply document-wide styling
        self._apply_styling(document)

        # Render document title if present
        if document.title:
            self._render_title(document.title)

        # Render all sections
        for section in document.sections:
            self._render_section(section)

        # Generate output path if not provided
        if output_path is None:
            temp_dir = Path(tempfile.gettempdir())
            filename = f"{document.document_type}_{uuid.uuid4().hex[:8]}.docx"
            output_path = temp_dir / filename

        # Ensure parent directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Save document
        self.doc.save(str(output_path))

        # WINDOWS FIX: Ensure file is fully written to disk
        if hasattr(os, 'fsync'):
            try:
                with open(output_path, 'rb') as f:
                    os.fsync(f.fileno())
                logger.debug(f"OS fsync completed for {output_path}")
            except Exception as e:
                logger.warning(f"OS fsync failed: {e}")

        # Wait for file system to stabilize (Windows-specific)
        time.sleep(0.05)  # 50ms delay

        # Validate ZIP structure (DOCX files are ZIP archives)
        try:
            with zipfile.ZipFile(output_path, 'r') as z:
                # Test ZIP integrity
                corrupt_file = z.testzip()
                if corrupt_file:
                    logger.error(f"Corrupted file in DOCX ZIP: {corrupt_file}")
                    raise ValueError(f"Corrupted file in DOCX ZIP: {corrupt_file}")

                # Verify required DOCX parts exist
                required_files = ['[Content_Types].xml', 'word/document.xml']
                zip_files = set(z.namelist())
                missing = [f for f in required_files if f not in zip_files]
                if missing:
                    logger.error(f"Missing required DOCX files: {missing}")
                    raise ValueError(f"Missing required DOCX files: {missing}")

                logger.info(f"ZIP validation passed: {len(zip_files)} files in archive")

        except zipfile.BadZipFile as e:
            logger.error(f"Invalid DOCX ZIP structure: {e}")
            raise ValueError(f"Generated DOCX file is corrupted: {e}")
        except Exception as e:
            logger.error(f"ZIP validation failed: {e}")
            raise

        logger.info(f"Document saved and validated: {output_path}")

        return output_path

    def _render_title(self, title: str) -> None:
        """
        Render the document title.

        Args:
            title: Document title text
        """
        # Add title as a centered heading
        para = self.doc.add_heading(title, level=0)
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _apply_styling(self, document: UniversalDocument) -> None:
        """
        Apply document-wide styling including fonts, page setup, headers/footers.

        Args:
            document: UIF document with styling configuration
        """
        # Apply document styling (fonts, sizes, colors)
        if document.styling:
            self.style_engine.apply_document_styling(document.styling.model_dump())

        # Setup page layout (size, margins)
        if document.page_setup:
            self.style_engine.setup_page_layout(document.page_setup.model_dump())

        # Add header/footer
        if document.header_footer:
            hf = document.header_footer

            if hf.header_text:
                self.style_engine.add_header(hf.header_text)

            if hf.footer_text or hf.show_page_numbers:
                footer_text = hf.footer_text or ""
                self.style_engine.add_footer(
                    footer_text,
                    show_page_number=hf.show_page_numbers,
                    page_number_position=hf.page_number_position,
                    include_total_pages=hf.include_total_pages,
                )

    def _render_section(self, section: Section) -> None:
        """
        Render a section with all its content.

        Args:
            section: Section to render
        """
        # Add section heading
        self.section_builder.add_heading(section.heading, section.level)

        # Render content blocks
        for block in section.content_blocks:
            self._render_content_block(block)

        # Render subsections recursively
        for subsection in section.subsections:
            self._render_section(subsection)

    def _render_content_block(self, block: ContentBlock) -> None:
        """
        Render a single content block based on its type.

        Args:
            block: Content block to render
        """
        try:
            if block.type == ContentBlockType.PARAGRAPH:
                self._render_paragraph(block)

            elif block.type == ContentBlockType.HEADING:
                level = block.level if block.level else 1
                self.section_builder.add_heading(block.content, level)

            elif block.type == ContentBlockType.PAGE_BREAK:
                self.section_builder.add_page_break()

            elif block.type == ContentBlockType.BULLET_LIST:
                self._render_bullet_list(block)

            elif block.type == ContentBlockType.NUMBERED_LIST:
                self._render_numbered_list(block)

            elif block.type == ContentBlockType.TABLE:
                self._render_table(block)

            elif block.type == ContentBlockType.SIGNATURE_BLOCK:
                self._render_signature_block(block)

            else:
                logger.warning(f"Unknown content block type: {block.type}")

        except Exception as e:
            logger.error(f"Error rendering content block of type {block.type}: {e}")
            # Continue rendering other blocks instead of failing completely
            raise

    def _render_paragraph(self, block: ContentBlock) -> None:
        """
        Render a paragraph content block.

        Args:
            block: Paragraph content block
        """
        # Convert to dict format expected by SectionBuilder
        block_dict = {
            "type": "paragraph",
            "content": block.content or "",
            "alignment": block.alignment.value if block.alignment else "left",
            "spacing_before": block.spacing_before,
            "spacing_after": block.spacing_after,
        }

        # Add formatting if present
        if block.formatting and block.formatting.ranges:
            block_dict["formatting"] = {
                "ranges": [
                    {
                        "start": r.start,
                        "end": r.end,
                        "bold": r.bold,
                        "italic": r.italic,
                        "underline": r.underline,
                    }
                    for r in block.formatting.ranges
                ]
            }

        self.section_builder.add_paragraph(block_dict)

    def _render_bullet_list(self, block: ContentBlock) -> None:
        """
        Render a bullet list content block.

        Args:
            block: Bullet list content block
        """
        if not block.items:
            logger.warning("Bullet list block has no items")
            return

        # Convert items to format expected by ListBuilder
        items = self._prepare_list_items(block.items)
        self.list_builder.add_bullet_list(items)

    def _render_numbered_list(self, block: ContentBlock) -> None:
        """
        Render a numbered list content block.

        Args:
            block: Numbered list content block
        """
        if not block.items:
            logger.warning("Numbered list block has no items")
            return

        # Determine style
        style = "decimal"
        if block.list_style:
            style = block.list_style.value

        # Convert items to format expected by ListBuilder
        items = self._prepare_list_items(block.items)
        self.list_builder.add_numbered_list(items, style=style)

    def _prepare_list_items(
        self, items: List[Any]
    ) -> List[Any]:
        """
        Prepare list items for the ListBuilder.

        Args:
            items: Raw list items (strings or dicts)

        Returns:
            Processed list items
        """
        processed = []
        for item in items:
            if isinstance(item, str):
                processed.append(item)
            elif isinstance(item, dict):
                # Pass through dict format
                processed.append(item)
            else:
                # Convert to string as fallback
                processed.append(str(item))
        return processed

    def _render_table(self, block: ContentBlock) -> None:
        """
        Render a table content block.

        Args:
            block: Table content block
        """
        if not block.table:
            logger.warning("Table block has no table data")
            return

        table_data = block.table

        # Convert to dict format expected by TableBuilder
        table_dict = {
            "type": "table",
            "headers": table_data.headers,
            "rows": self._prepare_table_rows(table_data.rows),
            "style": table_data.style,
            "header_background": table_data.header_background,
        }

        if table_data.column_widths:
            table_dict["column_widths"] = table_data.column_widths

        self.table_builder.add_table(table_dict)

    def _prepare_table_rows(
        self, rows: List[List[Any]]
    ) -> List[List[Any]]:
        """
        Prepare table rows for the TableBuilder.

        Args:
            rows: Raw table rows

        Returns:
            Processed table rows
        """
        processed_rows = []
        for row in rows:
            processed_row = []
            for cell in row:
                if isinstance(cell, str):
                    processed_row.append(cell)
                elif hasattr(cell, "model_dump"):
                    # Pydantic model (TableCell)
                    processed_row.append(cell.model_dump())
                elif isinstance(cell, dict):
                    processed_row.append(cell)
                else:
                    processed_row.append(str(cell))
            processed_rows.append(processed_row)
        return processed_rows

    def _render_signature_block(self, block: ContentBlock) -> None:
        """
        Render a signature block for consent forms.

        Args:
            block: Signature block content block
        """
        if not block.signature:
            logger.warning("Signature block has no signature data")
            return

        signature = block.signature

        # Add preamble if present
        if signature.preamble:
            para = self.doc.add_paragraph()
            para.add_run(signature.preamble)

        # Add signature lines
        for line in signature.lines:
            para = self.doc.add_paragraph()

            # Add label
            label_run = para.add_run(f"{line.label}: ")
            label_run.bold = True

            # Add signature line (underscores)
            line_length = int(line.line_width_inches * 20)
            para.add_run("_" * line_length)

            # Add some spacing after each signature line
            from docx.shared import Pt

            para.paragraph_format.space_after = Pt(12)

    async def render_and_upload(
        self,
        document: UniversalDocument,
        user_id: str,
        protocol_id: str,
        version: int,
        validate_first: bool = True,
    ) -> DocEngineResult:
        """
        Render document and upload to storage.

        Args:
            document: UIF document to render
            user_id: User ID for storage path
            protocol_id: Protocol ID for storage path
            version: Document version number
            validate_first: Whether to validate before rendering

        Returns:
            DocEngineResult with file path and metadata
        """
        # Render to temp file
        temp_path = self.render(document, validate_first=validate_first)

        try:
            # Read file data
            with open(temp_path, "rb") as f:
                file_data = f.read()

            file_size = len(file_data)

            # Generate filename
            doc_type = document.document_type.lower()
            filename = f"{doc_type}_v{version}.docx"

            # Upload to storage
            storage_path = await storage.upload_file(
                file_data=file_data,
                file_name=filename,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                folder=f"documents/{user_id}/{protocol_id}",
            )

            logger.info(f"Document uploaded to: {storage_path}")

            return DocEngineResult(
                file_path=storage_path,
                file_size=file_size,
                metadata={
                    "document_type": document.document_type,
                    "title": document.title,
                    "version": version,
                    "uif_version": document.uif_version,
                    "generated_by": document.compliance.generated_by,
                    "polished_by": document.compliance.polished_by,
                    "generated_at": document.compliance.generated_at.isoformat()
                    if document.compliance.generated_at
                    else None,
                    "section_count": document.count_sections(),
                    "content_block_count": document.count_content_blocks(),
                },
            )

        finally:
            # Cleanup temp file
            try:
                temp_path.unlink(missing_ok=True)
            except Exception as e:
                logger.warning(f"Failed to cleanup temp file {temp_path}: {e}")

    def render_to_bytes(
        self,
        document: UniversalDocument,
        validate_first: bool = True,
    ) -> bytes:
        """
        Render document directly to bytes without saving to disk.

        Args:
            document: UIF document to render
            validate_first: Whether to validate before rendering

        Returns:
            Document as bytes
        """
        import io

        logger.info(f"Rendering document to bytes: {document.title}")

        # Optionally validate first
        if validate_first:
            is_valid, errors = self.validate(document)
            if not is_valid:
                error_msg = "Document validation failed:\n" + "\n".join(
                    f"  - {e}" for e in errors
                )
                logger.error(error_msg)
                raise ValueError(error_msg)

        # Initialize fresh document
        self._init_document()

        # Apply document-wide styling
        self._apply_styling(document)

        # Render document title if present
        if document.title:
            self._render_title(document.title)

        # Render all sections
        for section in document.sections:
            self._render_section(section)

        # Save to BytesIO buffer
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)

        return buffer.read()


# Singleton instance
doc_engine = DocEngine()
